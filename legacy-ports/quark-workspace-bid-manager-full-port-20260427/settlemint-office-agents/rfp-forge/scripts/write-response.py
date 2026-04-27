#!/usr/bin/env python3
"""Write RFP responses back into the original document format.
Input: JSON file with responses (from the agent) + original document
Output: Modified document with responses filled in

Usage: python3 write-response.py <original_file> <responses_json> <output_file>

responses_json format:
{
  "responses": [
    {"question": "...", "response": "...", "confidence": "verified|docs|review", "row": 5, "col": 2, "sheet": "Sheet1"},
    ...
  ]
}
"""
import json, sys, os, re, shutil
from pathlib import Path

def write_xlsx(original, responses, output):
    """Write responses into Excel spreadsheet."""
    import openpyxl
    shutil.copy2(original, output)
    wb = openpyxl.load_workbook(output)
    
    for r in responses:
        sheet_name = r.get("sheet")
        row = r.get("row")
        response_col = r.get("response_col")
        
        if not sheet_name or not row:
            continue
        
        ws = wb[sheet_name]
        
        # If no response column was identified, find or create one
        if response_col is None:
            # Look for a "Response" header
            q_col = r.get("col", 0)
            header_row = r.get("header_row", 1)
            # Put response in the column after the question
            response_col = q_col + 1
            # Check if there's a header we should respect
            header_cell = ws.cell(row=header_row, column=response_col + 1)
            if not header_cell.value:
                header_cell.value = "SettleMint Response"
                header_cell.font = openpyxl.styles.Font(bold=True)
        
        # Write the response
        cell = ws.cell(row=row, column=response_col + 1)  # openpyxl is 1-indexed
        
        # Format: response + confidence indicator
        confidence = r.get("confidence", "")
        conf_marker = {"verified": "✅", "docs": "⚠️", "review": "❓"}.get(confidence, "")
        
        cell.value = r["response"]
        cell.alignment = openpyxl.styles.Alignment(wrap_text=True, vertical='top')
        
        # Add confidence in adjacent column if space
        if conf_marker:
            conf_cell = ws.cell(row=row, column=response_col + 2)
            if not conf_cell.value:
                conf_cell.value = conf_marker
    
    # Apply SettleMint branding to response cells
    try:
        sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "shared" / "scripts"))
        from excel_brand import get_brand_font, get_thin_border
        brand_font = get_brand_font(size=11)
        brand_border = get_thin_border()
        for r in responses:
            sheet_name = r.get("sheet")
            row = r.get("row")
            response_col = r.get("response_col", r.get("col", 0) + 1)
            if sheet_name and row:
                cell = wb[sheet_name].cell(row=row, column=response_col + 1)
                cell.font = brand_font
                cell.border = brand_border
    except ImportError:
        pass  # Graceful fallback if shared scripts not available

    wb.save(output)
    return {"ok": True, "output": output, "responses_written": len(responses)}


def write_docx(original, responses, output):
    """Write responses into Word document."""
    import docx
    shutil.copy2(original, output)
    doc = docx.Document(output)
    
    # Build lookup by table position and paragraph position
    table_responses = [r for r in responses if "table" in r]
    para_responses = [r for r in responses if "paragraph" in r]
    
    # Fill in table responses
    for r in table_responses:
        t_idx = r.get("table", 0)
        row_idx = r.get("row", 0)
        r_col = r.get("response_col", r.get("col", 0) + 1)
        
        if t_idx < len(doc.tables):
            table = doc.tables[t_idx]
            if row_idx < len(table.rows):
                row = table.rows[row_idx]
                # Ensure enough columns
                if r_col < len(row.cells):
                    cell = row.cells[r_col]
                    confidence = r.get("confidence", "")
                    conf_marker = {"verified": " ✅", "docs": " ⚠️", "review": " ❓"}.get(confidence, "")
                    cell.text = r["response"] + conf_marker
    
    # Insert responses after paragraph questions
    # Process in reverse order to preserve indices
    for r in sorted(para_responses, key=lambda x: x.get("paragraph", 0), reverse=True):
        p_idx = r.get("paragraph", 0)
        if p_idx < len(doc.paragraphs):
            # Insert a new paragraph after the question
            para = doc.paragraphs[p_idx]
            confidence = r.get("confidence", "")
            conf_marker = {"verified": " ✅", "docs": " ⚠️", "review": " ❓"}.get(confidence, "")
            
            # Add response paragraph after the question
            new_para = doc.add_paragraph()
            # Move it to after the question paragraph
            para._element.addnext(new_para._element)
            
            run = new_para.add_run(r["response"] + conf_marker)
            run.font.color.rgb = docx.shared.RGBColor(0, 0, 128)  # Dark blue for responses
            run.font.size = docx.shared.Pt(10)
    
    doc.save(output)
    return {"ok": True, "output": output, "responses_written": len(responses)}


def write_docx_from_pdf(responses, output):
    """Create a new Word document from PDF-extracted questions (PDF can't be written to)."""
    import docx
    doc = docx.Document()
    
    # Title
    doc.add_heading("RFP Response — SettleMint", level=1)
    doc.add_paragraph("Generated by Quark • Confidence: ✅ Verified in code  ⚠️ Based on docs  ❓ Needs review")
    doc.add_paragraph("")
    
    current_page = None
    for r in responses:
        page = r.get("page")
        if page and page != current_page:
            doc.add_heading(f"Page {page}", level=2)
            current_page = page
        
        # Question
        q_para = doc.add_paragraph()
        q_run = q_para.add_run(f"{r.get('id', '')} {r['question']}".strip())
        q_run.bold = True
        
        # Response
        confidence = r.get("confidence", "")
        conf_marker = {"verified": " ✅", "docs": " ⚠️", "review": " ❓"}.get(confidence, "")
        
        r_para = doc.add_paragraph()
        r_run = r_para.add_run(r["response"] + conf_marker)
        r_run.font.color.rgb = docx.shared.RGBColor(0, 0, 128)
        
        doc.add_paragraph("")  # Spacer
    
    doc.save(output)
    return {"ok": True, "output": output, "responses_written": len(responses)}


def main():
    print(json.dumps({
        "error": "Deprecated legacy output path: write-response.py is disabled. Canonical RFP Forge deliverables must be generated inside rfp-forge/output/{institution-slug}/ via the current procurement-generation workflow."
    }))
    sys.exit(1)


if __name__ == "__main__":
    main()
