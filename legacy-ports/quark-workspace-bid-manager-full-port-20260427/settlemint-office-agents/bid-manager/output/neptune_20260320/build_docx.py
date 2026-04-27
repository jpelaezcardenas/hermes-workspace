#!/usr/bin/env python3
"""Build Neptune executive summary DOCX from markdown + SettleMint master template."""

import re
import os
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_BREAK

TEMPLATE = '/Users/quark/Public/quark/workspace/settlemint-office-agents/bid-manager/templates/MASTER_TEMPLATE_LOCKED.docx'
MD_FILE = '/Users/quark/Public/quark/workspace/settlemint-office-agents/bid-manager/output/neptune_20260320/neptune-executive-summary.md'
OUTPUT = '/Users/quark/Public/quark/workspace/settlemint-office-agents/bid-manager/output/neptune_20260320/neptune-executive-summary.docx'

doc = Document(TEMPLATE)

# Remove dummy content (paragraphs 20 onward) — keep front-matter (0-19)
# We delete from the end to avoid index shifting
body = doc.element.body
paragraphs_to_remove = []
for i, p in enumerate(doc.paragraphs):
    if i >= 20:
        paragraphs_to_remove.append(p._element)

for elem in paragraphs_to_remove:
    body.remove(elem)

# Add a page break after front-matter
page_break_para = doc.add_paragraph()
page_break_para.style = doc.styles['Normal']
run = page_break_para.add_run()
run.add_break(WD_BREAK.PAGE)

# Parse markdown and add content
with open(MD_FILE, 'r') as f:
    lines = f.read().strip().split('\n')

i = 0
while i < len(lines):
    line = lines[i]

    # H1: # Executive Summary
    if line.startswith('# ') and not line.startswith('## '):
        title = line[2:].strip()
        p = doc.add_paragraph(title, style='H1')
        i += 1
        continue

    # H2: ## Section Title
    if line.startswith('## '):
        title = line[3:].strip()
        p = doc.add_paragraph(title, style='H2')
        i += 1
        continue

    # Empty line — skip
    if line.strip() == '':
        i += 1
        continue

    # Normal paragraph text — collect continuation lines
    para_text = line.strip()
    i += 1
    while i < len(lines) and lines[i].strip() != '' and not lines[i].startswith('#'):
        para_text += ' ' + lines[i].strip()
        i += 1

    p = doc.add_paragraph(para_text, style='Normal')
    # Ensure no Courier-like fonts sneak in — clear run-level font overrides
    for run in p.runs:
        rPr = run._element.find(qn('w:rPr'))
        if rPr is not None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                rPr.remove(rFonts)

doc.save(OUTPUT)

size = os.path.getsize(OUTPUT)
print(f"DOCX saved: {OUTPUT}")
print(f"File size: {size:,} bytes ({size/1024:.1f} KB)")
