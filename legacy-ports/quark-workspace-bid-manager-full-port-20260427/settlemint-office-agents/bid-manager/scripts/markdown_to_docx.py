#!/usr/bin/env python3
"""Convert Markdown to DOCX using python-docx.

Handles: H1-H4, paragraphs, bold, italic, bold-italic, bullet lists,
numbered lists, simple pipe-delimited markdown tables, page breaks (---),
and Mermaid diagrams (rendered to PNG/SVG via mmdc then embedded).

Prefers Figtree font if available, falls back to Calibri.
Headings use #000099 (dark navy). Table headers use navy background + white text.

Usage:
    python3 scripts/markdown_to_docx.py input.md [output.docx]
"""

import json
import os
import re
import subprocess
import sys
import tempfile
from pathlib import Path
from urllib.parse import unquote

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SHARED_SCRIPTS_DIR = Path(__file__).resolve().parents[2] / "shared" / "scripts"


def _strip_emoji(text: str) -> str:
    """Remove emoji characters from text for client-facing output.

    Strips Unicode emoji (including confidence dots, status indicators,
    and any other colored/graphical Unicode symbols) per the
    'No Emoji in Output Documents' directive (Gyan, 2026-04-03).
    """
    # Covers most emoji ranges: emoticons, dingbats, symbols, transport,
    # miscellaneous, flags, supplemental symbols, and variation selectors.
    emoji_pattern = re.compile(
        "["
        "\U0001F600-\U0001F64F"  # emoticons
        "\U0001F300-\U0001F5FF"  # symbols & pictographs
        "\U0001F680-\U0001F6FF"  # transport & map
        "\U0001F1E0-\U0001F1FF"  # flags
        "\U0001F900-\U0001F9FF"  # supplemental symbols
        "\U0001FA00-\U0001FA6F"  # chess symbols
        "\U0001FA70-\U0001FAFF"  # symbols extended-A
        "\U00002702-\U000027B0"  # dingbats
        "\U000024C2-\U0001F251"  # enclosed characters
        "\U0000FE0F"             # variation selector
        "\U0000200D"             # zero-width joiner
        "\U00002600-\U000026FF"  # misc symbols (includes ⚠️⛔)
        "\U00002B50-\U00002B55"  # stars
        "]+",
        flags=re.UNICODE,
    )
    return emoji_pattern.sub("", text)
if str(SHARED_SCRIPTS_DIR) not in sys.path:
    sys.path.insert(0, str(SHARED_SCRIPTS_DIR))

from office_output_validation import OfficeValidationError, validate_docx_against_template
from proposal_format_validator import TemplateFormatError, get_vague_error, preflight_template_check
from office_runtime_guard import append_build_metadata, assert_output_within_agent

MMDC_PATH = os.environ.get("MMDC_PATH", "mmdc")
MERMAID_CONFIG = {
    "theme": "base",
    "themeVariables": {
        "primaryColor": "#E8EAF6",
        "primaryTextColor": "#000099",
        "primaryBorderColor": "#000099",
        "lineColor": "#000099",
        "secondaryColor": "#FFF3E0",
        "tertiaryColor": "#E8F5E9",
        "background": "#FFFFFF",
        "fontFamily": "Figtree, Calibri, sans-serif",
        "fontSize": "14px"
    },
    "flowchart": {
        "curve": "basis",
        "padding": 50,
        "nodeSpacing": 70,
        "rankSpacing": 70,
        "htmlLabels": False,
        "diagramPadding": 20,
        "useMaxWidth": False,
        "wrappingWidth": 300
    },
    "sequence": {
        "mirrorActors": False,
        "diagramMarginX": 40,
        "diagramMarginY": 20,
        "messageMargin": 28,
        "boxMargin": 14,
        "actorMargin": 50,
        "width": 180,
        "height": 64
    }
}

# Page content area constraints (A4 with typical margins)
# Used for diagram sizing validation
MAX_CONTENT_WIDTH_INCHES = 6.0   # ~160mm content area
MAX_CONTENT_HEIGHT_INCHES = 9.0  # ~230mm between header and footer


PREFERRED_FONT = "Figtree"
FALLBACK_FONT = "Figtree"  # Figtree everywhere, per Gyan directive (2026-03-13)
FONT_SIZE = Pt(11)

# Gyan's canonical heading spec (updated 2026-03-13):
# ALL headings use Figtree. Figtree everywhere, no exceptions.
# H1: Figtree 16pt Bold #000099 | H2: Figtree 14pt Bold #000099
# H3: Figtree 12pt Bold #000099 | H4: Figtree 11pt Bold #000099
HEADING_SPEC = {
    1: {"font": "Figtree", "size": Pt(16), "style": "H1"},
    2: {"font": "Figtree", "size": Pt(14), "style": "H2"},
    3: {"font": "Figtree", "size": Pt(12), "style": "H3"},
    4: {"font": "Figtree", "size": Pt(11), "style": "H4"},
}
HEADING_COLOR = RGBColor(0x00, 0x00, 0x99)  # #000099 dark navy blue
TABLE_HEADER_BG = "E8E0F0"       # Light lavender header fill
TABLE_HEADER_TEXT = "2D1A8C"     # Dark purple/indigo header text
TABLE_BORDER_COLOR = "C4B4D8"    # Medium lavender border color
TABLE_FONT_SIZE = Pt(10)         # 10pt font for table cell text


def _pick_font() -> str:
    """Return Figtree if the font file exists locally, else Calibri."""
    figtree_dir = Path(__file__).resolve().parents[1] / ".." / "assets" / "fonts" / "figtree"
    if not figtree_dir.exists():
        figtree_dir = Path.home() / "Public" / "quark" / "workspace" / "assets" / "fonts" / "figtree"
    if figtree_dir.exists() and any(figtree_dir.glob("*.ttf")):
        return PREFERRED_FONT
    return FALLBACK_FONT


def _apply_font(run, font_name: str, size=None):
    run.font.name = font_name
    if size:
        run.font.size = size


def _add_rich_text(paragraph, text: str, font_name: str, font_size=None):
    """Parse inline markdown (bold, italic, bold-italic) into runs."""
    # Pattern order matters: bold-italic before bold before italic
    patterns = [
        (r"\*\*\*(.+?)\*\*\*", True, True),   # bold-italic
        (r"\*\*(.+?)\*\*", True, False),        # bold
        (r"\*(.+?)\*", False, True),             # italic
        (r"_(.+?)_", False, True),               # italic (underscore)
    ]

    segments: list[tuple[str, bool, bool]] = []
    remaining = text

    while remaining:
        earliest_match = None
        earliest_pos = len(remaining)
        matched_pattern = None

        for pattern, bold, italic in patterns:
            m = re.search(pattern, remaining)
            if m and m.start() < earliest_pos:
                earliest_match = m
                earliest_pos = m.start()
                matched_pattern = (bold, italic)

        if earliest_match and matched_pattern:
            # Text before the match
            if earliest_pos > 0:
                segments.append((remaining[:earliest_pos], False, False))
            segments.append((earliest_match.group(1), matched_pattern[0], matched_pattern[1]))
            remaining = remaining[earliest_match.end():]
        else:
            segments.append((remaining, False, False))
            remaining = ""

    for seg_text, bold, italic in segments:
        run = paragraph.add_run(seg_text)
        run.bold = bold
        run.italic = italic
        _apply_font(run, font_name, font_size or FONT_SIZE)


def _parse_table_line(line: str) -> list[str]:
    """Split a pipe-delimited markdown table row into cells."""
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def _is_separator_row(line: str) -> bool:
    """Check if a table line is a separator (e.g., |---|---|)."""
    return bool(re.match(r"^\|?[\s\-:|]+\|?$", line.strip()))


PUPPETEER_CONFIG = str(Path(__file__).resolve().parent / "puppeteer-config.json")


def _render_mermaid_to_image(code: str, output_dir: Path, index: int) -> Path | None:
    """Render Mermaid code to PNG (216 DPI via scale=3). Returns image path or None."""
    import time
    # Delay between renders to avoid Chrome resource exhaustion on macOS
    if index > 1:
        time.sleep(5)

    with tempfile.NamedTemporaryFile(mode="w", suffix=".mmd", delete=False) as f:
        f.write(code)
        mmd_path = f.name

    with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False) as f:
        json.dump(MERMAID_CONFIG, f)
        config_path = f.name

    png_path = output_dir / f"_diagram-{index:03d}.png"
    cmd = [MMDC_PATH, "-i", mmd_path, "-o", str(png_path),
           "--outputFormat", "png", "--scale", "3",
           "--configFile", config_path, "--backgroundColor", "white"]
    if os.path.exists(PUPPETEER_CONFIG):
        cmd.extend(["--puppeteerConfigFile", PUPPETEER_CONFIG])
    # Load Figtree font embedded as base64 woff2 for Puppeteer rendering
    figtree_css = Path(__file__).resolve().parents[2] / "shared" / "brand" / "figtree-embedded.css"
    if figtree_css.exists():
        cmd.extend(["--cssFile", str(figtree_css)])

    # Retry up to 3 times with increasing delay (Chrome can be flaky on macOS)
    for attempt in range(3):
        try:
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
            if result.returncode == 0 and png_path.exists():
                os.unlink(mmd_path)
                os.unlink(config_path)
                return png_path
            if attempt < 2:
                time.sleep(5 * (attempt + 1))
        except (subprocess.TimeoutExpired, FileNotFoundError) as e:
            if attempt < 2:
                time.sleep(5 * (attempt + 1))
            else:
                print(f"  Warning: Mermaid render error after 3 attempts: {e}")
                os.unlink(mmd_path)
                os.unlink(config_path)
                return None

    print(f"  Warning: Mermaid diagram {index} render failed after 3 attempts")
    os.unlink(mmd_path)
    os.unlink(config_path)
    return None


def _get_image_dimensions_inches(img_path: Path, render_scale: int = 3) -> tuple[float, float]:
    """Get image dimensions in inches, accounting for render scale.
    
    Mermaid renders at 3× scale (216 DPI), so divide pixel dimensions
    by (96 * scale) to get actual intended inches.
    """
    try:
        from PIL import Image
        with Image.open(img_path) as img:
            w_px, h_px = img.size
            dpi = 96 * render_scale  # 288 DPI for scale=3
            return w_px / dpi, h_px / dpi
    except ImportError:
        # Fallback: assume 6 inches wide (we'll cap at max anyway)
        return MAX_CONTENT_WIDTH_INCHES, MAX_CONTENT_HEIGHT_INCHES


def _fit_image_to_page(img_path: Path, render_scale: int = 3) -> Inches:
    """Calculate the optimal width for an image to fit within page content area.
    
    Returns the width as an Inches value that ensures:
    - Image does not exceed MAX_CONTENT_WIDTH_INCHES
    - Image does not exceed MAX_CONTENT_HEIGHT_INCHES (proportional)
    - Aspect ratio is preserved
    """
    w_in, h_in = _get_image_dimensions_inches(img_path, render_scale)
    
    # Scale to fit width
    if w_in > MAX_CONTENT_WIDTH_INCHES:
        scale = MAX_CONTENT_WIDTH_INCHES / w_in
        w_in *= scale
        h_in *= scale
    
    # Scale to fit height
    if h_in > MAX_CONTENT_HEIGHT_INCHES:
        scale = MAX_CONTENT_HEIGHT_INCHES / h_in
        w_in *= scale
        h_in *= scale
    
    return Inches(w_in)


TEMPLATE_PATH = Path(__file__).resolve().parents[1] / "templates" / "MASTER_TEMPLATE_LOCKED.docx"

# Native Word list numIds (defined in template numbering.xml)
# numId=1 = heading outline (DO NOT USE for lists)
# numId=4 = bullet list (abstractNumId=3)
# numId=5 = numbered list (abstractNumId=4)
BULLET_NUM_ID = "4"
NUMBERED_NUM_ID = "5"
NUMBERED_ABSTRACT_ID = "4"


def _set_numPr(paragraph, num_id: str, ilvl: int = 0):
    """Set native Word numbering on a paragraph."""
    pPr = paragraph._element.get_or_add_pPr()
    numPr = pPr.makeelement(qn('w:numPr'), {})
    ilvl_elem = numPr.makeelement(qn('w:ilvl'), {qn('w:val'): str(ilvl)})
    numId_elem = numPr.makeelement(qn('w:numId'), {qn('w:val'): num_id})
    numPr.append(ilvl_elem)
    numPr.append(numId_elem)
    pPr.append(numPr)


def _new_num_instance(doc, abstract_num_id: str) -> str:
    """Create a new numId instance pointing to an abstractNum, so numbering restarts.
    
    Each separate numbered list group needs its own numId to restart at 1.
    """
    from lxml import etree
    numbering = doc.part.numbering_part._element
    
    # Find max existing numId
    max_id = max(int(n.get(qn('w:numId'))) for n in numbering.findall(qn('w:num')))
    new_id = str(max_id + 1)
    
    # Create new num element
    num_elem = etree.SubElement(numbering, qn('w:num'))
    num_elem.set(qn('w:numId'), new_id)
    abs_ref = etree.SubElement(num_elem, qn('w:abstractNumId'))
    abs_ref.set(qn('w:val'), abstract_num_id)
    
    # Add lvlOverride to restart at 1
    lvl_override = etree.SubElement(num_elem, qn('w:lvlOverride'))
    lvl_override.set(qn('w:ilvl'), '0')
    start_override = etree.SubElement(lvl_override, qn('w:startOverride'))
    start_override.set(qn('w:val'), '1')
    
    return new_id


# Cover page placeholder replacements
COVER_REPLACEMENTS = {
    "[DD Month YYYY]": "valid_until",
    "[Full Name]": "contact_name",
    "[Job Title]": "contact_title",
    "[email@example.com]": "contact_email",
    "[Country Code - Number]": "contact_phone",
    "[the Company]": "company",
}


def _update_cover_page(doc, fields: dict):
    """Replace placeholder text in the cover page.
    
    Handles two types of placeholders:
    1. Structured Document Tags (sdt) with alias: Company, Title, Subtitle
    2. Regular paragraph text: [DD Month YYYY], [Full Name], etc.
    
    fields dict keys: company, title, subtitle, valid_until, contact_name,
                      contact_title, contact_email, contact_phone
    """
    from lxml import etree
    body = doc.element.body
    
    # Map sdt aliases to field keys
    SDT_FIELDS = {
        "Company": "company",
        "Title": "title",
        "Subtitle": "subtitle",
    }
    
    # Update structured document tags (content controls)
    for sdt in body.findall('.//' + qn('w:sdt')):
        sdtPr = sdt.find(qn('w:sdtPr'))
        if sdtPr is None:
            continue
        alias = sdtPr.find(qn('w:alias'))
        alias_val = alias.get(qn('w:val')) if alias is not None else ''
        
        if alias_val in SDT_FIELDS:
            field_key = SDT_FIELDS[alias_val]
            if field_key not in fields:
                continue
            
            value = fields[field_key]
            
            # Remove the showingPlcHdr flag (tells Word it's no longer placeholder)
            showing = sdtPr.find(qn('w:showingPlcHdr'))
            if showing is not None:
                sdtPr.remove(showing)
            
            # Remove the placeholder element entirely (some Word versions use
            # this to override display even after showingPlcHdr is removed)
            placeholder = sdtPr.find(qn('w:placeholder'))
            if placeholder is not None:
                sdtPr.remove(placeholder)
            
            # Remove dataBinding — Word overrides inline content with bound
            # document properties (which may be empty). We set doc properties
            # separately below, but removing the binding ensures inline content
            # is always shown.
            data_binding = sdtPr.find(qn('w:dataBinding'))
            if data_binding is not None:
                sdtPr.remove(data_binding)
            
            # Replace content: find paragraphs (may be direct or inside table cells)
            content = sdt.find(qn('w:sdtContent'))
            if content is not None:
                # Find first paragraph, whether direct child or inside a table cell
                first_p = content.find(qn('w:p'))
                if first_p is None:
                    # Look inside table cells
                    first_p = content.find('.//' + qn('w:p'))
                
                if first_p is not None:
                    # Get existing run properties for formatting
                    existing_rPr = None
                    for r in first_p.findall(qn('w:r')):
                        if existing_rPr is None:
                            rPr_elem = r.find(qn('w:rPr'))
                            if rPr_elem is not None:
                                existing_rPr = etree.tostring(rPr_elem)
                    
                    # Remove all runs
                    for r in first_p.findall(qn('w:r')):
                        first_p.remove(r)
                    
                    # Add new run with value
                    new_run = etree.SubElement(first_p, qn('w:r'))
                    if existing_rPr is not None:
                        new_run.append(etree.fromstring(existing_rPr))
                    new_t = etree.SubElement(new_run, qn('w:t'))
                    new_t.text = value
                    new_t.set(qn('xml:space'), 'preserve')
    
    # Also set document core/extended properties (belt-and-suspenders:
    # even though we removed dataBinding, Word may still check these)
    if 'title' in fields:
        doc.core_properties.title = fields['title']
    if 'subtitle' in fields:
        doc.core_properties.subject = fields['subtitle']
    # Extended properties (Company) need XML manipulation
    if 'company' in fields:
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            # Find the extended properties part (app.xml)
            for rel in doc.part.package.rels.values():
                if 'extended-properties' in rel.reltype:
                    app_part = rel.target_part
                    app_xml = app_part._element
                    ns = 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'
                    company_elem = app_xml.find(f'{{{ns}}}Company')
                    if company_elem is None:
                        company_elem = etree.SubElement(app_xml, f'{{{ns}}}Company')
                    company_elem.text = fields['company']
                    break
        except Exception:
            pass  # Not critical; dataBinding removal is the primary fix
    
    # Update regular paragraph placeholders (handles cross-run splits)
    for p in doc.paragraphs[:20]:
        full_text = p.text
        for placeholder, field_key in COVER_REPLACEMENTS.items():
            if placeholder in full_text and field_key in fields:
                _replace_across_runs(p, placeholder, fields[field_key])


def _replace_across_runs(paragraph, old_text: str, new_text: str):
    """Replace text that may span multiple runs in a paragraph."""
    runs = paragraph.runs
    if not runs:
        return
    
    # Build a map of character positions to runs
    full = ""
    run_map = []  # (run_index, char_within_run) for each char in full text
    for ri, r in enumerate(runs):
        for ci in range(len(r.text)):
            run_map.append((ri, ci))
        full += r.text
    
    pos = full.find(old_text)
    if pos < 0:
        return
    
    end_pos = pos + len(old_text)
    
    # Find which runs are affected
    start_run, start_char = run_map[pos]
    end_run, end_char = run_map[end_pos - 1]
    
    # Put replacement text in the first affected run, clear text from middle/end runs
    r0 = runs[start_run]
    r0.text = r0.text[:start_char] + new_text + runs[end_run].text[end_char + 1:]
    
    # Clear intermediate and end runs (if different from start)
    for ri in range(start_run + 1, end_run + 1):
        runs[ri].text = ""


def _add_toc(doc, insert_before_elem=None):
    """Update the template's existing TOC SDT with real heading entries.
    
    Finds the TOC SDT that already exists in the template, removes its
    placeholder content, and populates it with actual heading text from
    the document. Also adds a proper TOC field so Word can refresh it.
    
    If no TOC SDT exists, creates one from scratch (fallback).
    """
    from docx.oxml import OxmlElement
    from lxml import etree

    body = doc.element.body

    # Collect all headings from the document
    headings = []
    for para in doc.paragraphs:
        if para.style and para.style.name:
            sname = para.style.name
            hlevel = None
            if sname.startswith('Heading '):
                try:
                    hlevel = int(sname.replace('Heading ', '').strip())
                except ValueError:
                    pass
            elif sname in ('H1', 'H2', 'H3', 'H4'):
                hlevel = int(sname[1])
            if hlevel and 1 <= hlevel <= 4 and para.text.strip():
                headings.append((hlevel, para.text.strip()))

    # Find existing TOC SDT in the document
    # Detection: check for TOC field instruction OR docPartGallery="Table of Contents"
    toc_sdt = None
    for sdt in body.findall('.//' + qn('w:sdt')):
        sdtPr = sdt.find(qn('w:sdtPr'))
        sdtContent = sdt.find(qn('w:sdtContent'))
        # Check docPartGallery (template-style TOC)
        if sdtPr is not None:
            dpo = sdtPr.find(qn('w:docPartObj'))
            if dpo is not None:
                gallery = dpo.find(qn('w:docPartGallery'))
                if gallery is not None and 'Table of Contents' in gallery.get(qn('w:val'), ''):
                    toc_sdt = sdt
                    break
        # Check for TOC field instruction
        if sdtContent is not None:
            for instr in sdtContent.findall('.//' + qn('w:instrText')):
                if instr.text and 'TOC' in instr.text:
                    toc_sdt = sdt
                    break
            if toc_sdt is not None:
                break

    if toc_sdt is not None:
        # Update existing TOC SDT - find its content container
        sdtContent = toc_sdt.find(qn('w:sdtContent'))
        if sdtContent is None:
            sdtContent = toc_sdt

        # Remove all existing paragraphs in the TOC content
        for p in list(sdtContent.findall(qn('w:p'))):
            sdtContent.remove(p)

        # --- TOC Title: insert BEFORE the SDT as a regular paragraph ---
        # This keeps it outside Word's TOC field so our formatting sticks.
        body = doc.element.body
        title_p = OxmlElement('w:p')
        # Paragraph spacing: 0 before, 240 twips (12pt) after
        title_pPr = etree.SubElement(title_p, qn('w:pPr'))
        title_spacing = etree.SubElement(title_pPr, qn('w:spacing'))
        title_spacing.set(qn('w:before'), '0')
        title_spacing.set(qn('w:after'), '240')
        title_r = etree.SubElement(title_p, qn('w:r'))
        title_rPr = etree.SubElement(title_r, qn('w:rPr'))
        title_bold = etree.SubElement(title_rPr, qn('w:b'))
        title_sz = etree.SubElement(title_rPr, qn('w:sz'))
        title_sz.set(qn('w:val'), '32')  # 16pt
        title_szCs = etree.SubElement(title_rPr, qn('w:szCs'))
        title_szCs.set(qn('w:val'), '32')
        title_color = etree.SubElement(title_rPr, qn('w:color'))
        title_color.set(qn('w:val'), '000072')  # Theme blue matching headings
        title_font = etree.SubElement(title_rPr, qn('w:rFonts'))
        title_font.set(qn('w:ascii'), 'Figtree')
        title_font.set(qn('w:hAnsi'), 'Figtree')
        title_font.set(qn('w:cs'), 'Figtree')
        title_t = etree.SubElement(title_r, qn('w:t'))
        title_t.text = "Table of Contents"
        title_t.set(qn('xml:space'), 'preserve')
        # Insert title paragraph right before the TOC SDT
        body.insert(list(body).index(toc_sdt), title_p)

        # --- TOC field inside SDT ---
        toc_field_p = etree.SubElement(sdtContent, qn('w:p'))
        toc_field_r = etree.SubElement(toc_field_p, qn('w:r'))
        fldChar_begin = etree.SubElement(toc_field_r, qn('w:fldChar'))
        fldChar_begin.set(qn('w:fldCharType'), 'begin')
        instrText = etree.SubElement(toc_field_r, qn('w:instrText'))
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = r' TOC \o "1-4" \h \z \u '
        fldChar_separate = etree.SubElement(toc_field_r, qn('w:fldChar'))
        fldChar_separate.set(qn('w:fldCharType'), 'separate')

        # --- TOC entries with page numbers and dotted leaders ---
        # Right tab stop at 9072 twips (~6.3 inches / ~16cm) with dot leader
        RIGHT_TAB_POS = '9072'
        for entry_idx, (level, text) in enumerate(headings):
            entry_p = etree.SubElement(sdtContent, qn('w:p'))
            pPr = etree.SubElement(entry_p, qn('w:pPr'))
            # Use TOC style
            toc_style_id = f'TOC{level}'
            pStyle = etree.SubElement(pPr, qn('w:pStyle'))
            pStyle.set(qn('w:val'), toc_style_id)
            # Add right-aligned tab stop with dot leader
            tabs = etree.SubElement(pPr, qn('w:tabs'))
            tab = etree.SubElement(tabs, qn('w:tab'))
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:leader'), 'dot')
            tab.set(qn('w:pos'), RIGHT_TAB_POS)

            # Heading text
            entry_r = etree.SubElement(entry_p, qn('w:r'))
            entry_t = etree.SubElement(entry_r, qn('w:t'))
            entry_t.text = text
            entry_t.set(qn('xml:space'), 'preserve')
            # Tab character
            tab_r = etree.SubElement(entry_p, qn('w:r'))
            tab_elem = etree.SubElement(tab_r, qn('w:tab'))
            # Page number placeholder
            pn_r = etree.SubElement(entry_p, qn('w:r'))
            pn_t = etree.SubElement(pn_r, qn('w:t'))
            pn_t.text = str(entry_idx + 3)  # Approximate page numbers (cover=1, TOC=2, content starts ~3)
            pn_t.set(qn('xml:space'), 'preserve')

        # Close the TOC field
        end_p = etree.SubElement(sdtContent, qn('w:p'))
        end_r = etree.SubElement(end_p, qn('w:r'))
        fldChar_end = etree.SubElement(end_r, qn('w:fldChar'))
        fldChar_end.set(qn('w:fldCharType'), 'end')

        print(f"  TOC updated with {len(headings)} headings from template SDT (with TOC field)")
    else:
        # Fallback: no existing TOC SDT found, create one from scratch
        print("  Warning: No template TOC SDT found, creating new TOC")
        _add_toc_fresh(doc, insert_before_elem, headings)


def _add_toc_fresh(doc, insert_before_elem, headings):
    """Fallback: create a new TOC field from scratch (when no template TOC exists)."""
    from docx.oxml import OxmlElement
    from lxml import etree

    body = doc.element.body

    def _insert_or_append(element):
        if insert_before_elem is not None:
            body.insert(list(body).index(insert_before_elem), element)
        else:
            body.append(element)

    # TOC Title
    toc_heading = doc.add_paragraph()
    run = toc_heading.add_run("Table of Contents")
    run.bold = True
    run.font.name = "Figtree"
    run.font.size = Pt(16)
    run.font.color.rgb = HEADING_COLOR
    body.remove(toc_heading._element)
    _insert_or_append(toc_heading._element)

    # TOC field
    toc_para = doc.add_paragraph()
    run = toc_para.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r' TOC \o "1-4" \h \z \u '
    run._element.append(instrText)

    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar_separate)

    body.remove(toc_para._element)
    _insert_or_append(toc_para._element)

    last_entry_elem = toc_para._element

    if headings:
        for level, text in headings:
            entry_para = doc.add_paragraph()
            entry_para.style = doc.styles['Normal']
            indent_inches = (level - 1) * 0.25
            entry_para.paragraph_format.left_indent = Inches(indent_inches)
            entry_para.paragraph_format.space_before = Pt(2)
            entry_para.paragraph_format.space_after = Pt(2)
            entry_para.paragraph_format.tab_stops.add_tab_stop(
                Inches(6.0), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=2
            )

            entry_run = entry_para.add_run(text)
            entry_run.font.size = Pt(10)
            entry_run.font.name = "Figtree"

            body.remove(entry_para._element)
            _insert_or_append(entry_para._element)
            last_entry_elem = entry_para._element

    end_run_elem = OxmlElement('w:r')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    end_run_elem.append(fldChar_end)
    last_entry_elem.append(end_run_elem)

    # Page break
    pb_para = doc.add_paragraph()
    pb_run = pb_para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    pb_run._element.append(br)
    body.remove(pb_para._element)
    _insert_or_append(pb_para._element)


def _remove_empty_pages(doc):
    """Scan for and remove elements that would create blank pages.
    
    Checks for:
    - Consecutive page/section breaks with only empty paragraphs between them
    - Empty paragraphs left behind after a page/section break before the next heading
    - Page breaks immediately followed by headings that already have 'page break before'
    - Trailing empty paragraphs before sectPr
    """
    body = doc.element.body
    
    def _has_page_break(elem):
        """Check if a paragraph element contains an explicit page break."""
        for br in elem.findall('.//' + qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                return True
        return False
    
    def _has_section_break(elem):
        """Check if a paragraph element contains a section break."""
        pPr = elem.find(qn('w:pPr'))
        return pPr is not None and pPr.find(qn('w:sectPr')) is not None
    
    def _has_page_or_section_break(elem):
        return _has_page_break(elem) or _has_section_break(elem)
    
    def _is_empty_or_whitespace(elem):
        """Check if a paragraph has no text or only whitespace."""
        text = ''.join(t.text or '' for t in elem.findall('.//' + qn('w:t')))
        return not text.strip()
    
    def _has_page_break_before_style(elem):
        """Check if a paragraph's style includes 'page break before'."""
        pPr = elem.find(qn('w:pPr'))
        if pPr is not None:
            pb_before = pPr.find(qn('w:pageBreakBefore'))
            if pb_before is not None:
                val = pb_before.get(qn('w:val'), 'true')
                return val.lower() in ('true', '1', 'on', '')
        return False
    
    def _is_heading_paragraph(elem):
        """Check if a paragraph is one of the heading styles we generate/use."""
        pPr = elem.find(qn('w:pPr'))
        if pPr is None:
            return False
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is None:
            return False
        style_val = pStyle.get(qn('w:val'), '')
        return style_val in {'H1', 'H2', 'H3', 'H4', 'Heading1', 'Heading2', 'Heading3', 'Heading4'}
    
    def _is_toc_sdt(elem):
        """Check if an SDT element is the template TOC container."""
        if (elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag) != 'sdt':
            return False
        sdtPr = elem.find(qn('w:sdtPr'))
        sdtContent = elem.find(qn('w:sdtContent'))
        if sdtPr is not None:
            dpo = sdtPr.find(qn('w:docPartObj'))
            if dpo is not None:
                gallery = dpo.find(qn('w:docPartGallery'))
                if gallery is not None and 'Table of Contents' in gallery.get(qn('w:val'), ''):
                    return True
        if sdtContent is not None:
            for instr in sdtContent.findall('.//' + qn('w:instrText')):
                if instr.text and 'TOC' in instr.text:
                    return True
        return False
    
    to_remove = set()
    
    # Pass 1: Remove consecutive page/section breaks (with only empty paragraphs between)
    elements = list(body)
    prev_break_idx = None
    
    for idx, elem in enumerate(elements):
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag == 'sectPr':
            continue
        if tag != 'p':
            prev_break_idx = None
            continue
        
        has_break = _has_page_or_section_break(elem)
        is_empty = _is_empty_or_whitespace(elem)
        
        if has_break:
            if prev_break_idx is not None:
                all_between_empty = all(
                    elements[j].tag.split('}')[-1] == 'p' and _is_empty_or_whitespace(elements[j])
                    for j in range(prev_break_idx + 1, idx)
                    if elements[j].tag.split('}')[-1] != 'sectPr'
                )
                if all_between_empty:
                    to_remove.add(id(elem))
                    for j in range(prev_break_idx + 1, idx):
                        if elements[j].tag.split('}')[-1] == 'p' and _is_empty_or_whitespace(elements[j]):
                            to_remove.add(id(elements[j]))
            prev_break_idx = idx
        elif is_empty:
            pass
        else:
            prev_break_idx = None
    
    # Pass 2: After any page/section break (or TOC block), remove consecutive
    # empty paragraphs until the next heading. This catches the cover→TOC→content
    # transition where template empties can create a blank page.
    elements = list(body)
    seen_break_or_toc = False
    for elem in elements:
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag == 'sectPr':
            continue
        if tag == 'sdt':
            if _is_toc_sdt(elem):
                seen_break_or_toc = True
            continue
        if tag != 'p':
            seen_break_or_toc = False
            continue
        if _has_page_or_section_break(elem):
            seen_break_or_toc = True
            continue
        if not seen_break_or_toc:
            continue
        if _is_empty_or_whitespace(elem):
            to_remove.add(id(elem))
            continue
        if _is_heading_paragraph(elem) or _has_page_break_before_style(elem):
            seen_break_or_toc = False
            continue
        seen_break_or_toc = False
    
    # Pass 3: Remove page breaks immediately before headings with 'page break before'
    elements = list(body)
    for idx, elem in enumerate(elements):
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag != 'p' or not _has_page_break(elem):
            continue
        if id(elem) in to_remove:
            continue
        
        # Look ahead past empty paragraphs for a heading with page-break-before
        for j in range(idx + 1, len(elements)):
            next_tag = elements[j].tag.split('}')[-1] if '}' in elements[j].tag else elements[j].tag
            if next_tag == 'sectPr':
                break
            if next_tag != 'p':
                break
            if _is_empty_or_whitespace(elements[j]) and not _has_page_break_before_style(elements[j]):
                continue  # skip empty paragraphs
            if _has_page_break_before_style(elements[j]):
                # The heading already forces a page break — remove the explicit one
                # Only remove if the page break paragraph is otherwise empty
                if _is_empty_or_whitespace(elem):
                    to_remove.add(id(elem))
            break
    
    # Pass 4: Remove trailing empty/whitespace paragraphs before sectPr
    elements = list(body)
    sect_idx = None
    for idx, elem in enumerate(elements):
        if (elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag) == 'sectPr':
            sect_idx = idx
            break
    
    if sect_idx is not None:
        for j in range(sect_idx - 1, -1, -1):
            tag = elements[j].tag.split('}')[-1] if '}' in elements[j].tag else elements[j].tag
            if tag != 'p':
                break
            if _is_empty_or_whitespace(elements[j]) and not _has_page_break(elements[j]):
                to_remove.add(id(elements[j]))
            else:
                break
    
    # Apply removals
    removed = 0
    for elem in list(body):
        if id(elem) in to_remove:
            body.remove(elem)
            removed += 1
    
    if removed:
        print(f"  Empty page prevention: removed {removed} redundant element(s)")


def convert(input_path: str, output_path: str | None = None, cover_fields: dict | None = None) -> Path:
    src = Path(input_path)
    if not src.exists():
        raise FileNotFoundError(f"Input file not found: {src}")

    dst = Path(output_path) if output_path else src.with_suffix(".docx")
    dst = assert_output_within_agent("bid-manager", dst)
    font_name = _pick_font()

    # Pre-flight: verify locked template is present and unmodified before generation begins
    try:
        preflight_template_check(TEMPLATE_PATH)
    except TemplateFormatError as exc:
        raise OfficeValidationError(get_vague_error()) from exc

    doc = Document(str(TEMPLATE_PATH))
    body = doc.element.body

    # Template structure:
    #   [0] <sdt> — cover page SDT (structured doc tag)
    #   [1-14] Cover page paragraphs (confidentiality, T&C, contacts)
    #   [15] Page break
    #   [16] Empty paragraph
    #   [17] <sdt> — TOC SDT
    #   [18-21] Empty + page break
    #   [22+] Sample content (to be replaced)
    #   [-1] <sectPr> — section properties
    #
    # Strategy: keep cover page elements [0..21], remove sample content [22+],
    # preserve sectPr. New content appends after the cover/TOC.

    # Find all elements
    elements = list(body)
    # Identify the boundary: keep cover page (sdt + paragraphs up to first page break),
    # keep TOC (sdt after first page break), remove everything after second page break.
    # The boundary is the element with the second page break itself (inclusive).
    page_break_count = 0
    keep_until = -1  # index of last element to keep
    for idx, elem in enumerate(elements):
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag == 'sectPr':
            continue
        if tag == 'p':
            # Check for page break
            has_pb = False
            for br in elem.findall('.//' + qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    has_pb = True
            if has_pb:
                page_break_count += 1
            if page_break_count >= 2:
                # This paragraph with the second page break is the last we keep
                keep_until = idx
                break

    # Remove everything after keep_until, except sectPr
    to_remove = []
    for idx, elem in enumerate(elements):
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag == 'sectPr':
            continue
        if idx > keep_until and keep_until >= 0:
            to_remove.append(elem)

    for elem in to_remove:
        body.remove(elem)

    print(f"  Using template: {TEMPLATE_PATH.name} ({len(doc.styles)} styles)")
    print(f"  Cover page preserved ({keep_until + 1} elements), {len(to_remove)} sample elements removed")

    # Set default font on the document style
    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = FONT_SIZE

    # Update cover page placeholder fields if provided
    if cover_fields:
        _update_cover_page(doc, cover_fields)

    raw_text = src.read_text(encoding="utf-8")
    raw_text = _strip_emoji(raw_text)  # No emoji in output (Gyan directive, 2026-04-03)
    lines = raw_text.splitlines()
    i = 0
    diagram_count = 0
    in_bullet_list = False
    in_num_list = False
    current_num_id = NUMBERED_NUM_ID
    first_h1_skipped = False  # Skip leading H1s — they're document titles (already on cover page)
    first_content_emitted = False

    while i < len(lines):
        line = lines[i]

        # Blank line
        if not line.strip():
            i += 1
            continue

        # Page break (--- on its own line)
        if line.strip() == "---":
            # Mark that we've passed the title block — H1s after this are real sections
            first_h1_skipped = True
            if not first_content_emitted:
                i += 1
                continue  # Skip page breaks before first real content
            doc.add_page_break()
            i += 1
            continue

        # Mermaid code block
        if line.strip().startswith("```mermaid"):
            mermaid_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                mermaid_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # skip closing ```
            diagram_count += 1
            mermaid_code = "\n".join(mermaid_lines)
            img_path = _render_mermaid_to_image(mermaid_code, dst.parent, diagram_count)
            if img_path:
                fit_width = _fit_image_to_page(img_path)
                doc.add_picture(str(img_path), width=fit_width)
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # Fallback: add as italic text noting render failure
                p = doc.add_paragraph()
                run = p.add_run(f"[Diagram {diagram_count} — render failed, see markdown source]")
                run.italic = True
                _apply_font(run, font_name, FONT_SIZE)
            continue

        # Skip other code blocks (non-mermaid)
        if line.strip().startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                i += 1
            if i < len(lines):
                i += 1
            continue

        # Markdown images: ![alt text](path)
        img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", line.strip())
        if img_match:
            alt_text = img_match.group(1)
            img_rel_path = img_match.group(2).strip()

            # Markdown generators sometimes percent-encode spaces and other path
            # characters. Decode before resolving against the markdown source dir.
            candidate_rel_paths = [img_rel_path]
            decoded_rel_path = unquote(img_rel_path)
            if decoded_rel_path != img_rel_path:
                candidate_rel_paths.insert(0, decoded_rel_path)

            candidate_abs_paths = []
            for rel_path in candidate_rel_paths:
                candidate_abs_paths.append((src.parent / rel_path).resolve())
                # Fallback: try one directory higher (handles ../../shared/ when shared/
                # is a sibling of the parent project dir, not inside it)
                if rel_path.startswith("../../"):
                    alt_rel = rel_path.replace("../../", "../../../", 1)
                    candidate_abs_paths.append((src.parent / alt_rel).resolve())

            img_abs = None
            for candidate in candidate_abs_paths:
                candidate_jpg = candidate.with_suffix('.jpg')
                if candidate_jpg.exists():
                    img_abs = candidate_jpg
                    break
                if candidate.exists():
                    img_abs = candidate
                    break

            if img_abs:
                fit_width = _fit_image_to_page(img_abs, render_scale=1)
                doc.add_picture(str(img_abs), width=fit_width)
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Add caption below image if alt text exists
                if alt_text:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cap.add_run(alt_text)
                    run.italic = True
                    run.font.size = Pt(8)
                    _apply_font(run, font_name, Pt(8))
            else:
                # Image not found - add placeholder text
                p = doc.add_paragraph()
                run = p.add_run(f"[Image: {alt_text or img_rel_path}]")
                run.italic = True
                _apply_font(run, font_name, FONT_SIZE)
            i += 1
            continue

        # Headings
        heading_match = re.match(r"^(#{1,4})\s+(.+)$", line)
        if heading_match:
            level = len(heading_match.group(1))  # 1-4

            # Skip leading H1s that appear before the first page break (---).
            # These are document titles/subtitles already on the cover page.
            # H1s after the first --- are real content sections (e.g., Executive Summary).
            if level == 1 and not first_h1_skipped:
                i += 1
                continue

            # Promote heading levels by 1 for outline numbering since H1 is
            # consumed by the cover/title. Body sections use H2 in markdown but
            # should render as outline level 1 ("1.", not "1.1") in Word.
            # H2→H1 (ilvl 0), H3→H2 (ilvl 1), H4→H3 (ilvl 2).
            effective_level = max(1, level - 1)
            spec = HEADING_SPEC.get(effective_level, HEADING_SPEC[4])

            # Apply template style (H1/H2/H3/H4) and explicitly set outline
            # numbering. Style inheritance (H2 → Heading 2 → outline) is unreliable
            # in Word, so we set numPr directly on every heading paragraph.
            style_name = spec["style"]
            try:
                heading = doc.add_paragraph(style=style_name)
            except KeyError:
                heading = doc.add_heading("", level=effective_level)

            # Explicitly link to heading outline numbering (numId=1)
            # ilvl = effective_level - 1 (promoted H2→ilvl 0 = "1.", H3→ilvl 1 = "1.1", etc.)
            _set_numPr(heading, "1", effective_level - 1)

            heading.clear()
            # Strip leading section numbers (e.g. "1.", "1.1", "2.3.1.") since
            # Word outline numbering (numId=1) already provides heading numbers.
            heading_text = re.sub(r'^\d+(\.\d+)*\.?\s+', '', heading_match.group(2)).strip()
            heading.add_run(heading_text)
            # No explicit font/size/color/bold on the run — style handles it all
            first_content_emitted = True
            i += 1
            continue

        # Bullet list — native Word bullet with multi-level support
        # Level determined by leading whitespace: 0 spaces = level 0, 2-3 spaces = level 1, 4-5 = level 2
        bullet_match = re.match(r"^(\s*)([-*+])\s+(.+)$", line)
        if bullet_match:
            indent = len(bullet_match.group(1))
            level = min(indent // 2, 2)  # max 3 levels (0, 1, 2)
            if not in_bullet_list:
                in_bullet_list = True
            p = doc.add_paragraph(style="List Paragraph")
            _set_numPr(p, BULLET_NUM_ID, level)
            _add_rich_text(p, bullet_match.group(3), font_name)
            first_content_emitted = True
            i += 1
            continue
        else:
            in_bullet_list = False

        # Numbered list — native Word numbering with multi-level support
        # Each separate list group gets a fresh numId instance to restart at 1
        num_match = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
        if num_match:
            indent = len(num_match.group(1))
            level = min(indent // 2, 2)
            if not in_num_list:
                in_num_list = True
                current_num_id = _new_num_instance(doc, NUMBERED_ABSTRACT_ID)
            p = doc.add_paragraph(style="List Paragraph")
            _set_numPr(p, current_num_id, level)
            _add_rich_text(p, num_match.group(2), font_name)
            first_content_emitted = True
            i += 1
            continue
        else:
            in_num_list = False

        # Table detection: line starts with | and contains at least one more |
        if "|" in line and line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                if not _is_separator_row(lines[i]):
                    table_lines.append(_parse_table_line(lines[i]))
                i += 1

            if table_lines:
                first_content_emitted = True
                num_cols = max(len(row) for row in table_lines)
                table = doc.add_table(rows=len(table_lines), cols=num_cols)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Apply custom border color (lavender)
                tbl = table._tbl
                tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
                borders = OxmlElement('w:tblBorders')
                for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '4')  # 0.5pt
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), TABLE_BORDER_COLOR)
                    borders.append(border)
                # Remove existing borders if any
                existing_borders = tblPr.find(qn('w:tblBorders'))
                if existing_borders is not None:
                    tblPr.remove(existing_borders)
                tblPr.append(borders)

                # Cell margins — 0.04" (~58 twips) on all four sides
                existing_margins = tblPr.find(qn('w:tblCellMar'))
                if existing_margins is not None:
                    tblPr.remove(existing_margins)
                cell_mar = OxmlElement('w:tblCellMar')
                for side, val in [('top', '58'), ('bottom', '58'), ('left', '58'), ('right', '58')]:
                    margin = OxmlElement(f'w:{side}')
                    margin.set(qn('w:w'), val)
                    margin.set(qn('w:type'), 'dxa')
                    cell_mar.append(margin)
                tblPr.append(cell_mar)

                # Smart column widths: analyze content to determine optimal widths
                # Total available width ~9072 twips (~16cm / ~6.3 inches)
                TOTAL_TABLE_WIDTH = 9072
                col_max_lens = [0] * num_cols
                for row_data in table_lines:
                    for ci, ct in enumerate(row_data):
                        if ci < num_cols:
                            col_max_lens[ci] = max(col_max_lens[ci], len(ct))
                # Compute proportional widths based on content length, with min/max constraints
                total_len = sum(col_max_lens) or 1
                col_widths = []
                for ci in range(num_cols):
                    proportion = col_max_lens[ci] / total_len
                    # Min 8% per column, max 60% for any single column
                    proportion = max(0.08, min(0.60, proportion))
                    col_widths.append(proportion)
                # Normalize to sum to 1.0
                width_sum = sum(col_widths)
                col_widths = [w / width_sum for w in col_widths]
                # Apply widths to table grid
                tbl_grid = tbl.find(qn('w:tblGrid'))
                if tbl_grid is not None:
                    grid_cols = tbl_grid.findall(qn('w:gridCol'))
                    for ci, gc in enumerate(grid_cols):
                        if ci < num_cols:
                            gc.set(qn('w:w'), str(int(TOTAL_TABLE_WIDTH * col_widths[ci])))

                for row_idx, cells in enumerate(table_lines):
                    row = table.rows[row_idx]
                    # Set row heights
                    if row_idx == 0:
                        row.height = Cm(0.75)  # header row
                        row.height_rule = 1   # WD_ROW_HEIGHT_RULE.AT_LEAST
                    else:
                        row.height = Cm(0.75)  # data row
                        row.height_rule = 1   # WD_ROW_HEIGHT_RULE.AT_LEAST

                    for col_idx, cell_text in enumerate(cells):
                        if col_idx < num_cols:
                            cell = table.cell(row_idx, col_idx)
                            cell.text = ""
                            p = cell.paragraphs[0]
                            # Vertical alignment: center for all cells
                            cell.vertical_alignment = 1  # WD_ALIGN_VERTICAL.CENTER
                            # Cell paragraph spacing: 1.15x line spacing
                            p_pPr = p._element.get_or_add_pPr()
                            p_spacing = OxmlElement('w:spacing')
                            p_spacing.set(qn('w:line'), '276')   # 1.15x
                            p_spacing.set(qn('w:lineRule'), 'auto')
                            p_spacing.set(qn('w:before'), '0')
                            p_spacing.set(qn('w:after'), '0')
                            p_pPr.append(p_spacing)
                            _add_rich_text(p, cell_text, font_name, TABLE_FONT_SIZE)
                            # Style header row
                            if row_idx == 0:
                                for run in p.runs:
                                    run.bold = True
                                    run.font.color.rgb = RGBColor(0x2D, 0x1A, 0x8C)
                                # Alignment: first column left, all others center
                                if col_idx == 0:
                                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                else:
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                # Set cell background to lavender
                                shading = cell._element.get_or_add_tcPr()
                                shading_elm = shading.makeelement(qn('w:shd'), {
                                    qn('w:fill'): TABLE_HEADER_BG,
                                    qn('w:val'): 'clear',
                                })
                                shading.append(shading_elm)
                # Add spacing after table (12pt gap before next element)
                spacer = doc.add_paragraph()
                spacer_pPr = spacer._element.get_or_add_pPr()
                spacer_spacing = OxmlElement('w:spacing')
                spacer_spacing.set(qn('w:before'), '0')
                spacer_spacing.set(qn('w:after'), '0')
                spacer_spacing.set(qn('w:line'), '120')  # 6pt line height — small spacer
                spacer_spacing.set(qn('w:lineRule'), 'exact')
                spacer_pPr.append(spacer_spacing)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_rich_text(p, line, font_name)
        first_content_emitted = True
        i += 1

    # Find the insertion point: after the cover page (second page break)
    # The content starts after the TOC, which is after cover page
    elements = list(body)
    pb_count = 0
    toc_insert_before = None
    for idx, elem in enumerate(elements):
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag == 'sectPr':
            continue
        if tag == 'p':
            for br in elem.findall('.//' + qn('w:br')):
                if br.get(qn('w:type')) == 'page':
                    pb_count += 1
            if pb_count >= 2:
                # TOC should be inserted right after this element (end of cover page)
                if idx + 1 < len(elements):
                    toc_insert_before = elements[idx + 1]
                break

    # Update the template's existing TOC SDT with real heading entries
    _add_toc(doc, insert_before_elem=toc_insert_before)

    # Empty page prevention: scan and remove consecutive page breaks
    _remove_empty_pages(doc)

    doc.save(str(dst))

    validate_docx_against_template(
        dst,
        TEMPLATE_PATH,
        allowed_signature_overrides={
            # Generated output hashes differ from locked template when python-docx rewrites
            # styles.xml and numbering.xml during save (canonical output format).
            # These hashes are from verified proposal-bank outputs built against the
            # MASTER_TEMPLATE_LOCKED.docx (Mar 2026 version). All prior accepted builds
            # share these signatures.
            "word/styles.xml": {
                "bbd6942c9b6f8be4164105f170626af0428e2e2dc6068a09124c8abab7f652d5",
            },
            "word/numbering.xml": {
                "4eeb495bdeb132f43b1691c6322e57b49d345a139642d52509d4618aca4b4614",
                # Additional numbering.xml variant from proposals with richer list structures
                "a80464e25fffcb701ec3325cfc7443f0c464923937af8523f60b0bd75662f5f3",
                # Variant from heading-number-strip fix (2026-04-03)
                "401bafa11f50168707b4420ca428deca0225ecdf4c6dd26dbcd62316dc6e0c20",
            },
        },
        required_document_tokens=("Company", "Title", "Subtitle", "TOC"),
    )
    append_build_metadata(
        "bid-manager",
        output_path=dst,
        operation="markdown_to_docx",
        extra={
            "input": str(src),
            "template": str(TEMPLATE_PATH),
            "cover_fields_supplied": bool(cover_fields),
        },
    )
    print(f"✓ Converted {src.name} → {dst.name} (locked template validation verified)")
    return dst


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print(__doc__.strip())
        sys.exit(1)

    # Parse optional --cover-json argument for cover page fields
    cover = None
    args = sys.argv[1:]
    if "--cover-json" in args:
        idx = args.index("--cover-json")
        cover = json.loads(args[idx + 1])
        args = args[:idx] + args[idx + 2:]

    try:
        convert(args[0], args[1] if len(args) > 1 else None, cover_fields=cover)
    except OfficeValidationError as err:
        print(get_vague_error(), file=sys.stderr)
        sys.exit(1)
