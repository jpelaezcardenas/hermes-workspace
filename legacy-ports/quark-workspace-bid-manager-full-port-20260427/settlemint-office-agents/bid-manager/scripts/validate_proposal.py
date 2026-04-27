#!/usr/bin/env python3
"""
Validate a bid-manager proposal for completeness.

Checks markdown proposals for:
- Minimum mermaid diagram count per variant
- Presence of image references
- Emoji contamination (forbidden in output)
- Section structure completeness
- Mandatory diagram sections coverage

Checks DOCX proposals for:
- Embedded image count
- Basic structure validation

Usage:
    python3 validate_proposal.py <file.md> [variant]
    python3 validate_proposal.py <file.docx> [variant]
    python3 validate_proposal.py <file.md> --variant full --strict
    python3 validate_proposal.py <file.md> --json

Variants: full (default), medium, compact
"""
import sys
import re
import json
import argparse
import unicodedata
from pathlib import Path


# Minimum mermaid block counts per variant
MINIMUMS = {
    "full": 30,
    "medium": 20,
    "compact": 10,
}

# Minimum screenshot (image reference) counts per variant
SCREENSHOT_MINIMUMS = {
    "full": 12,
    "medium": 8,
    "compact": 6,
}

# Minimum screenshot section spread per variant
SCREENSHOT_SECTION_MINIMUMS = {
    "full": 4,
    "medium": 3,
    "compact": 2,
}

# Minimum screenshot category variety per variant
SCREENSHOT_CATEGORY_MINIMUMS = {
    "full": 4,
    "medium": 3,
    "compact": 2,
}

# Minimum Mermaid sequence diagram counts per variant
SEQUENCE_DIAGRAM_MINIMUMS = {
    "full": 8,
    "medium": 6,
    "compact": 4,
}

# Mandatory diagram sections for full variant (section heading substring -> diagram type)
MANDATORY_DIAGRAM_SECTIONS_FULL = {
    "Platform Overview": "platform architecture layers",
    "Lifecycle": "lifecycle state",
    "Compliance": "compliance flow",
    "Security": "security layers or defense-in-depth",
    "Integration": "integration architecture",
    "Deployment": "deployment topology",
    "Implementation": "implementation timeline or gantt",
    "Access Control": "RBAC or access control model",
    "Data Architecture": "data flow",
    "Key Management": "custody or key management flow",
}

MANDATORY_DIAGRAM_SECTIONS_MEDIUM = {
    "Platform Overview": "platform architecture layers",
    "Lifecycle": "lifecycle state",
    "Compliance": "compliance flow",
    "Security": "security layers",
    "Integration": "integration architecture",
    "Deployment": "deployment topology",
    "Implementation": "implementation timeline",
}

MANDATORY_DIAGRAM_SECTIONS_COMPACT = {
    "Platform": "platform architecture",
    "Solution": "solution architecture",
    "Compliance": "compliance flow",
}


def count_mermaid_blocks(content: str) -> int:
    """Count mermaid code blocks in markdown content."""
    return len(re.findall(r"```mermaid", content))


def count_image_refs(content: str) -> int:
    """Count markdown image references."""
    return len(re.findall(r"!\[.*?\]\(.*?\)", content))


def extract_image_refs(content: str) -> list:
    """Extract markdown image references with surrounding section context."""
    lines = content.splitlines()
    current_section = "(no heading found)"
    refs = []

    for i, line in enumerate(lines):
        heading_match = re.match(r"^#{1,6}\s+(.+)$", line.strip())
        if heading_match:
            current_section = heading_match.group(1).strip()

        image_match = re.search(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            caption_line = lines[i + 1].strip() if i + 1 < len(lines) else ""
            refs.append({
                "alt": image_match.group(1).strip(),
                "path": image_match.group(2).strip(),
                "section": current_section,
                "line": i + 1,
                "has_figure_caption": bool(re.match(r"^\*Figure\s+\d+:", caption_line)),
                "caption_line": caption_line,
            })
    return refs


def screenshot_category_from_path(path: str) -> str:
    """Infer screenshot category from a shared screenshot path."""
    match = re.search(r"dalp-screenshots/([^/]+)/", path)
    if match:
        return match.group(1)
    match = re.search(r"new-2026-03-24/([^/]+)/", path)
    if match:
        return f"new-2026-03-24/{match.group(1)}"
    return "unknown"


def analyze_screenshot_policy(content: str, variant: str) -> dict:
    """Check caption coverage, section spread, and category variety for screenshots."""
    refs = extract_image_refs(content)
    section_count = len({ref['section'] for ref in refs})
    category_count = len({screenshot_category_from_path(ref['path']) for ref in refs})
    missing_captions = [
        {
            "line": ref["line"],
            "section": ref["section"],
            "path": ref["path"],
        }
        for ref in refs
        if not ref["has_figure_caption"]
    ]

    min_sections = SCREENSHOT_SECTION_MINIMUMS.get(variant, 2)
    min_categories = SCREENSHOT_CATEGORY_MINIMUMS.get(variant, 2)

    return {
        "count": len(refs),
        "section_count": section_count,
        "section_minimum": min_sections,
        "section_passed": section_count >= min_sections,
        "category_count": category_count,
        "category_minimum": min_categories,
        "category_passed": category_count >= min_categories,
        "missing_captions": missing_captions,
        "caption_passed": len(missing_captions) == 0,
    }


def count_sequence_diagrams(content: str) -> int:
    """Count Mermaid sequence diagram blocks."""
    blocks = re.findall(r"```mermaid\s*\n\s*sequenceDiagram", content)
    return len(blocks)


def _is_emoji(char: str) -> bool:
    """Check if a character is an emoji (not a standard punctuation/modifier like backtick)."""
    cp = ord(char)
    # Common emoji ranges
    if cp in range(0x1F600, 0x1F650):  # Emoticons
        return True
    if cp in range(0x1F300, 0x1F5FF):  # Misc Symbols and Pictographs
        return True
    if cp in range(0x1F680, 0x1F6FF):  # Transport and Map
        return True
    if cp in range(0x1F900, 0x1F9FF):  # Supplemental Symbols
        return True
    if cp in range(0x2600, 0x26FF):  # Misc Symbols
        return True
    if cp in range(0x2700, 0x27BF):  # Dingbats
        return True
    if cp in range(0xFE00, 0xFE0F):  # Variation Selectors
        return True
    if cp in range(0x1FA00, 0x1FA6F):  # Chess, extended-A
        return True
    if cp in range(0x1FA70, 0x1FAFF):  # Symbols extended-A
        return True
    if cp in range(0x200D, 0x200E):  # ZWJ
        return True
    # Specific common emoji
    if cp in (0x2B50, 0x2B55, 0x2B06, 0x2B07, 0x2B05, 0x27A1,  # stars, arrows
              0x2705, 0x274C, 0x274E, 0x2757, 0x2753, 0x2755,  # check, cross, exclamation
              0x26A0, 0x26D4, 0x2934, 0x2935):  # warning, no entry
        return True
    # Red/green/yellow/white circles and squares (the confidence tag dots)
    if cp in (0x1F534, 0x1F535, 0x1F7E0, 0x1F7E1, 0x1F7E2, 0x1F7E3,
              0x1F7E4, 0x1F7E5, 0x1F7E6, 0x1F7E7, 0x1F7E8, 0x1F7E9,
              0x26AA, 0x26AB, 0x2B1B, 0x2B1C):
        return True
    return False


def count_emoji(content: str) -> int:
    """Count emoji characters that should not appear in output."""
    count = 0
    for char in content:
        if _is_emoji(char):
            count += 1
    return count


def find_emoji_positions(content: str, max_report: int = 10) -> list:
    """Find positions and characters of emoji in the content."""
    positions = []
    for i, char in enumerate(content):
        if _is_emoji(char):
            # Find surrounding context
            start = max(0, i - 20)
            end = min(len(content), i + 20)
            context = content[start:end].replace("\n", " ")
            positions.append({
                "char": char,
                "position": i,
                "unicode": f"U+{ord(char):04X}",
                "context": f"...{context}...",
            })
            if len(positions) >= max_report:
                break
    return positions


def extract_sections_with_mermaid(content: str) -> dict:
    """Map H2/H3 section headings to whether they contain mermaid blocks."""
    sections = {}
    # Split by headings
    parts = re.split(r"(^#{2,3}\s+.+$)", content, flags=re.MULTILINE)

    current_heading = None
    current_body = ""

    for part in parts:
        heading_match = re.match(r"^#{2,3}\s+(.+)$", part.strip())
        if heading_match:
            # Save previous section
            if current_heading:
                sections[current_heading] = "```mermaid" in current_body
            current_heading = heading_match.group(1).strip()
            current_body = ""
        else:
            current_body += part

    # Save last section
    if current_heading:
        sections[current_heading] = "```mermaid" in current_body

    return sections


def check_mandatory_diagram_sections(content: str, variant: str) -> list:
    """Check that mandatory diagram sections contain mermaid blocks."""
    section_map = extract_sections_with_mermaid(content)

    if variant == "full":
        mandatory = MANDATORY_DIAGRAM_SECTIONS_FULL
    elif variant == "medium":
        mandatory = MANDATORY_DIAGRAM_SECTIONS_MEDIUM
    else:
        mandatory = MANDATORY_DIAGRAM_SECTIONS_COMPACT

    missing = []
    for section_keyword, diagram_type in mandatory.items():
        found_section = False
        has_diagram = False
        for heading, has_mermaid in section_map.items():
            if section_keyword.lower() in heading.lower():
                found_section = True
                if has_mermaid:
                    has_diagram = True
                    break
        if found_section and not has_diagram:
            missing.append({
                "section": section_keyword,
                "expected_diagram": diagram_type,
                "status": "section exists but no mermaid block found",
            })
        elif not found_section:
            missing.append({
                "section": section_keyword,
                "expected_diagram": diagram_type,
                "status": "section not found in document",
            })

    return missing


def check_banned_phrases(content: str) -> list:
    """Check for banned positioning phrases and em dashes."""
    banned = [
        ("\u2014", "Em dash (\u2014)", "Use comma, period, colon, or restructure the sentence"),
        ("production-grade", "production-grade", "State the specific capability instead"),
        ("production grade", "production grade", "State the specific capability instead"),
        ("enterprise-grade", "enterprise-grade", "State the specific capability instead"),
        ("enterprise grade", "enterprise grade", "State the specific capability instead"),
        ("pilot to production", "pilot to production", "Say what the platform does, not what journey it enables"),
        ("from ambition to", "from ambition to", "Start with the point, not buildup"),
        ("institutional-grade", "institutional-grade", "State the specific institutional requirement instead"),
    ]
    findings = []
    for pattern, label, fix_hint in banned:
        occurrences = []
        lines = content.split("\n")
        for i, line in enumerate(lines, 1):
            if pattern in line.lower() if pattern not in ("\u2014",) else pattern in line:
                occurrences.append({
                    "line": i,
                    "context": line.strip()[:120],
                })
        if occurrences:
            findings.append({
                "phrase": label,
                "count": len(occurrences),
                "fix_hint": fix_hint,
                "occurrences": occurrences[:5],
            })
    return findings


def check_placeholder_remnants(content: str) -> list:
    """Check for unresolved placeholder patterns."""
    patterns = [
        (r"\[VARIABLE:\s*[^\]]+\]", "Unresolved [VARIABLE: ...] placeholder"),
        (r"\[REUSABLE BLOCK:\s*[^\]]+\]", "Unexpanded [REUSABLE BLOCK: ...] tag"),
        (r"\[TODO[:\s]*[^\]]*\]", "Unresolved [TODO] marker"),
    ]
    findings = []
    for pattern, description in patterns:
        matches = re.findall(pattern, content)
        if matches:
            findings.append({
                "type": description,
                "count": len(matches),
                "examples": matches[:3],
            })
    return findings


def validate_markdown(md_path: str, variant: str = "full", strict: bool = False) -> dict:
    """Validate proposal markdown meets requirements."""
    content = Path(md_path).read_text(encoding="utf-8")

    mermaid_count = count_mermaid_blocks(content)
    min_required = MINIMUMS.get(variant, 10)
    image_count = count_image_refs(content)
    emoji_count = count_emoji(content)
    emoji_positions = find_emoji_positions(content) if emoji_count > 0 else []
    word_count = len(content.split())
    missing_diagrams = check_mandatory_diagram_sections(content, variant)
    placeholder_remnants = check_placeholder_remnants(content)
    banned_phrase_findings = check_banned_phrases(content)

    screenshot_min = SCREENSHOT_MINIMUMS.get(variant, 3)
    screenshot_policy = analyze_screenshot_policy(content, variant)
    screenshot_pass = image_count >= screenshot_min
    screenshot_captions_pass = screenshot_policy["caption_passed"]
    screenshot_sections_pass = screenshot_policy["section_passed"]
    screenshot_categories_pass = screenshot_policy["category_passed"]

    seq_diagram_count = count_sequence_diagrams(content)
    seq_diagram_min = SEQUENCE_DIAGRAM_MINIMUMS.get(variant, 1)
    seq_diagram_pass = seq_diagram_count >= seq_diagram_min

    mermaid_pass = mermaid_count >= min_required
    emoji_pass = emoji_count == 0
    diagram_sections_pass = len(missing_diagrams) == 0
    placeholders_pass = len(placeholder_remnants) == 0 if strict else True

    # Overall pass/fail
    passed = (
        mermaid_pass
        and emoji_pass
        and screenshot_pass
        and screenshot_captions_pass
        and screenshot_sections_pass
        and screenshot_categories_pass
        and seq_diagram_pass
    )
    if strict:
        passed = passed and diagram_sections_pass and placeholders_pass

    results = {
        "file": md_path,
        "format": "markdown",
        "variant": variant,
        "strict": strict,
        "passed": passed,
        "checks": {
            "mermaid_blocks": {
                "count": mermaid_count,
                "minimum": min_required,
                "passed": mermaid_pass,
            },
            "image_refs": {
                "count": image_count,
                "minimum": screenshot_min,
                "passed": screenshot_pass,
            },
            "screenshot_policy": {
                "caption_passed": screenshot_captions_pass,
                "missing_captions": screenshot_policy["missing_captions"],
                "section_count": screenshot_policy["section_count"],
                "section_minimum": screenshot_policy["section_minimum"],
                "section_passed": screenshot_sections_pass,
                "category_count": screenshot_policy["category_count"],
                "category_minimum": screenshot_policy["category_minimum"],
                "category_passed": screenshot_categories_pass,
            },
            "sequence_diagrams": {
                "count": seq_diagram_count,
                "minimum": seq_diagram_min,
                "passed": seq_diagram_pass,
            },
            "emoji": {
                "count": emoji_count,
                "passed": emoji_pass,
                "positions": emoji_positions,
            },
            "word_count": word_count,
            "mandatory_diagram_sections": {
                "missing": missing_diagrams,
                "passed": diagram_sections_pass,
            },
            "placeholders": {
                "remnants": placeholder_remnants,
                "passed": placeholders_pass,
            },
            "banned_phrases": {
                "findings": banned_phrase_findings,
                "count": sum(f["count"] for f in banned_phrase_findings),
                "passed": len(banned_phrase_findings) == 0,
            },
        },
    }

    return results


def validate_docx(docx_path: str, variant: str = "full") -> dict:
    """Validate proposal DOCX for embedded images and structure."""
    try:
        from docx import Document
    except ImportError:
        return {
            "file": docx_path,
            "format": "docx",
            "variant": variant,
            "passed": False,
            "error": "python-docx not installed. Run: pip install python-docx",
        }

    try:
        doc = Document(docx_path)
    except Exception as e:
        return {
            "file": docx_path,
            "format": "docx",
            "variant": variant,
            "passed": False,
            "error": f"Failed to open DOCX: {e}",
        }

    min_required = MINIMUMS.get(variant, 10)

    # Count inline images
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1

    # Count paragraphs and headings
    heading_count = 0
    paragraph_count = 0
    table_count = len(doc.tables)
    for para in doc.paragraphs:
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            heading_count += 1
        elif para.text.strip():
            paragraph_count += 1

    # Check for TOC field
    has_toc = False
    for para in doc.paragraphs:
        for run in para.runs:
            if run._element.xml and "TOC" in run._element.xml:
                has_toc = True
                break
        if has_toc:
            break

    image_pass = image_count >= min_required

    results = {
        "file": docx_path,
        "format": "docx",
        "variant": variant,
        "passed": image_pass,
        "checks": {
            "embedded_images": {
                "count": image_count,
                "minimum": min_required,
                "passed": image_pass,
            },
            "headings": heading_count,
            "paragraphs": paragraph_count,
            "tables": table_count,
            "has_toc": has_toc,
        },
    }

    return results


def print_results(results: dict) -> None:
    """Print human-readable validation results."""
    status = "PASS" if results["passed"] else "FAIL"
    print(f"\nProposal Validation: {status}")
    print(f"  File: {results['file']}")
    print(f"  Format: {results['format']}")
    print(f"  Variant: {results['variant']}")

    if "error" in results:
        print(f"  Error: {results['error']}")
        return

    checks = results["checks"]

    if results["format"] == "markdown":
        mc = checks["mermaid_blocks"]
        marker = "PASS" if mc["passed"] else "FAIL - BELOW MINIMUM"
        print(f"  Mermaid diagrams: {mc['count']}/{mc['minimum']} required [{marker}]")

        ic = checks["image_refs"]
        marker = "PASS" if ic["passed"] else "FAIL - BELOW MINIMUM"
        print(f"  Screenshots/images: {ic['count']}/{ic['minimum']} required [{marker}]")

        sp = checks["screenshot_policy"]
        marker = "PASS" if sp["caption_passed"] else "FAIL - MISSING FIGURE CAPTIONS"
        print(f"  Screenshot captions: [{marker}]")
        if sp["missing_captions"]:
            for item in sp["missing_captions"][:5]:
                print(f"    - L{item['line']} in {item['section']}: {item['path']}")

        marker = "PASS" if sp["section_passed"] else "FAIL - TOO FEW SECTIONS WITH SCREENSHOTS"
        print(f"  Screenshot section spread: {sp['section_count']}/{sp['section_minimum']} required [{marker}]")

        marker = "PASS" if sp["category_passed"] else "FAIL - TOO FEW SCREENSHOT CATEGORIES"
        print(f"  Screenshot category variety: {sp['category_count']}/{sp['category_minimum']} required [{marker}]")

        sd = checks["sequence_diagrams"]
        marker = "PASS" if sd["passed"] else "FAIL - BELOW MINIMUM"
        print(f"  Sequence diagrams: {sd['count']}/{sd['minimum']} required [{marker}]")

        print(f"  Word count: {checks['word_count']}")

        ec = checks["emoji"]
        marker = "PASS" if ec["passed"] else "FAIL - EMOJI FOUND"
        print(f"  Emoji characters: {ec['count']} [{marker}]")
        if ec["positions"]:
            for pos in ec["positions"][:5]:
                print(f"    {pos['char']} ({pos['unicode']}) at position {pos['position']}: {pos['context']}")

        mds = checks["mandatory_diagram_sections"]
        marker = "PASS" if mds["passed"] else "WARN - MISSING DIAGRAMS IN SECTIONS"
        print(f"  Mandatory diagram sections: {marker}")
        if mds["missing"]:
            for m in mds["missing"]:
                print(f"    - {m['section']}: {m['status']} (expected: {m['expected_diagram']})")

        ph = checks["placeholders"]
        if ph["remnants"]:
            print(f"  Placeholder remnants: {len(ph['remnants'])} types found")
            for r in ph["remnants"]:
                print(f"    - {r['type']}: {r['count']} instances (e.g. {r['examples'][0]})")

        bp = checks["banned_phrases"]
        marker = "PASS" if bp["passed"] else f"WARN - {bp['count']} BANNED PHRASE OCCURRENCES"
        print(f"  Banned phrases: {marker}")
        if bp["findings"]:
            for f in bp["findings"]:
                print(f"    - \"{f['phrase']}\": {f['count']} occurrences ({f['fix_hint']})")
                for occ in f["occurrences"][:2]:
                    print(f"      L{occ['line']}: {occ['context'][:80]}")

    elif results["format"] == "docx":
        ic = checks["embedded_images"]
        marker = "PASS" if ic["passed"] else "FAIL - BELOW MINIMUM"
        print(f"  Embedded images: {ic['count']}/{ic['minimum']} required [{marker}]")
        print(f"  Headings: {checks['headings']}")
        print(f"  Paragraphs: {checks['paragraphs']}")
        print(f"  Tables: {checks['tables']}")
        print(f"  Has TOC: {checks['has_toc']}")

    print()


def main():
    parser = argparse.ArgumentParser(
        description="Validate a bid-manager proposal for completeness."
    )
    parser.add_argument("file", help="Path to markdown (.md) or DOCX (.docx) file")
    parser.add_argument(
        "variant",
        nargs="?",
        default="full",
        choices=["full", "medium", "compact"],
        help="Proposal variant (default: full)",
    )
    parser.add_argument(
        "--strict",
        action="store_true",
        help="Strict mode: also fail on missing diagram sections and unresolved placeholders",
    )
    parser.add_argument(
        "--json",
        action="store_true",
        dest="json_output",
        help="Output results as JSON instead of human-readable text",
    )

    args = parser.parse_args()
    filepath = Path(args.file)

    if not filepath.exists():
        print(f"Error: File not found: {args.file}", file=sys.stderr)
        sys.exit(2)

    if filepath.suffix.lower() == ".docx":
        results = validate_docx(str(filepath), args.variant)
    elif filepath.suffix.lower() == ".md":
        results = validate_markdown(str(filepath), args.variant, args.strict)
    else:
        print(f"Error: Unsupported file type: {filepath.suffix}", file=sys.stderr)
        print("Supported formats: .md, .docx", file=sys.stderr)
        sys.exit(2)

    if args.json_output:
        print(json.dumps(results, indent=2, default=str))
    else:
        print_results(results)

    sys.exit(0 if results["passed"] else 1)


if __name__ == "__main__":
    main()
