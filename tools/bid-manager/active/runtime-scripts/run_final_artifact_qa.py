#!/usr/bin/env python3
"""Run final client-ready artifact QA regression for bid-manager."""
from __future__ import annotations

import argparse
import json
import re
import subprocess
import time
import zipfile
from pathlib import Path
from typing import Any

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
WRAPPER = ROOT / "bin" / "bid-manager"
BID = ROOT / "legacy" / "bid-manager"

BANNED_TEXT = [
    "proposal-maker",
    "proposal-factory",
    "buyer-safe",
    "client-ready",
    "the proposal should",
    "recommended next step",
    "internal review",
    "drafting note",
    "assembly note",
    "PLACEHOLDER",
    "[VARIABLE",
    "TODO",
]

FIXTURE_MD = """# Emirates NBD Tokenization Platform Response

## Executive Summary

Emirates NBD can use SettleMint's Digital Asset Lifecycle Platform to operate controlled EVM-based tokenized asset workflows with clear governance, auditable activity, and a pragmatic implementation path. The response focuses on production control: who can do what, how changes are approved, how lifecycle events are evidenced, and how the bank keeps custody and regulatory boundaries explicit.

## Proposed Operating Model

SettleMint provides the platform control plane for issuance, role-based permissions, lifecycle operations, and integration with the bank's surrounding systems. The bank remains the owner of participant policy, custody-provider selection, signatory authority, and final operating procedures. That boundary matters: it keeps the solution credible for bank risk, compliance, legal, and operations teams.

## Platform Capabilities

| Capability | Response |
|---|---|
| Tokenized asset lifecycle | Supports EVM-based issuance, transfer controls, servicing events, and operational oversight. |
| Governance | Supports maker-checker patterns, role-based permissions, approval evidence, and auditable operational activity. |
| Integration | Supports API-led integration patterns for custody, reporting, compliance, and enterprise workflow systems. |
| Deployment | Supports bank-grade implementation planning with environment separation, testing, cutover, and support rhythm. |

## Security and Control Evidence

The implementation should produce evidence that risk and audit teams can inspect: permission models, change records, transaction approvals, lifecycle event history, and exception handling. The platform response is designed to reduce ambiguity during procurement by separating native platform capability from bank-owned operating decisions.

## Implementation Approach

The recommended implementation starts with a controlled scope, validates the operating model with business and risk stakeholders, confirms custody and signing boundaries, and then expands into additional asset classes or lifecycle events. This avoids a fragile proof-of-concept and creates a reusable foundation for future tokenization work.

## Conclusion

SettleMint gives Emirates NBD a credible path to industrialize tokenized asset operations while preserving bank control over governance, custody boundaries, compliance evidence, and delivery risk.
"""


def run(name: str, cmd: list[str]) -> dict[str, Any]:
    start = time.time()
    proc = subprocess.run(cmd, text=True, capture_output=True)
    return {
        "name": name,
        "cmd": cmd,
        "exit": proc.returncode,
        "ok": proc.returncode == 0,
        "duration_seconds": round(time.time() - start, 3),
        "stdout_tail": proc.stdout[-3000:],
        "stderr_tail": proc.stderr[-3000:],
    }


def extract_docx_text(path: Path) -> tuple[str, int, int, int]:
    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            text += "\n" + " | ".join(cell.text for cell in row.cells)
    words = len(re.findall(r"\b\w+\b", text))
    return text, words, len(doc.paragraphs), len(doc.tables)


def docx_final_gate(path: Path) -> dict[str, Any]:
    result: dict[str, Any] = {"path": str(path), "exists": path.exists(), "ok": False}
    if not path.exists():
        return result
    with zipfile.ZipFile(path) as zf:
        bad = zf.testzip()
        names = zf.namelist()
    result["zip_ok"] = bad is None
    result["media_count"] = len([n for n in names if n.startswith("word/media/")])
    text, words, paragraphs, tables = extract_docx_text(path)
    lower = text.lower()
    hits = [term for term in BANNED_TEXT if term.lower() in lower]
    result.update({
        "word_count": words,
        "paragraphs": paragraphs,
        "tables": tables,
        "banned_text_hits": hits,
        "has_table_of_contents_text": "table of contents" in lower,
        "mentions_bid_manager": "bid-manager" in lower,
    })
    result["ok"] = bool(result["zip_ok"] and words >= 400 and not hits and not result["mentions_bid_manager"] and result["media_count"] >= 1)
    return result


def pdf_gate(path: Path) -> dict[str, Any]:
    result: dict[str, Any] = {"path": str(path), "exists": path.exists(), "ok": False}
    if not path.exists():
        return result
    info = run("pdfinfo", ["pdfinfo", str(path)])
    text = run("pdftotext", ["pdftotext", str(path), "-"])
    pages = None
    for line in info["stdout_tail"].splitlines():
        if line.startswith("Pages:"):
            pages = int(line.split(":", 1)[1].strip())
    result.update({"pdfinfo_ok": info["ok"], "pdftotext_ok": text["ok"], "pages": pages, "text_chars": len(text["stdout_tail"])})
    result["ok"] = bool(info["ok"] and text["ok"] and pages and pages >= 1 and result["text_chars"] > 300)
    return result


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Run final bid-manager artifact QA regression")
    parser.add_argument("--json", action="store_true")
    parser.add_argument("--out-dir", default=None)
    ns = parser.parse_args(argv)
    stamp = time.strftime("%Y%m%dT%H%M%SZ", time.gmtime())
    out_dir = Path(ns.out_dir).resolve() if ns.out_dir else BID / "output" / f"final-artifact-qa-{stamp}"
    out_dir.mkdir(parents=True, exist_ok=True)
    md = out_dir / "enbd-tokenization-final-qa.md"
    docx = out_dir / "enbd-tokenization-final-qa.docx"
    pdf = out_dir / "enbd-tokenization-final-qa.pdf"
    md.write_text(FIXTURE_MD, encoding="utf-8")
    cases = [run("md_to_docx", [str(WRAPPER), "md-to-docx", str(md), str(docx)])]
    if cases[-1]["ok"]:
        cases.append(run("render_pdf", ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx)]))
    gates = [docx_final_gate(docx), pdf_gate(pdf)]
    ok = all(c["ok"] for c in cases) and all(g["ok"] for g in gates)
    payload = {"agent": "bid-manager", "artifact_dir": str(out_dir), "cases": cases, "gates": gates, "ok": ok}
    (out_dir / "final-artifact-qa.json").write_text(json.dumps(payload, indent=2), encoding="utf-8")
    print(json.dumps(payload, indent=2) if ns.json else f"bid-manager final artifact QA: {'OK' if ok else 'FAIL'} artifacts={out_dir}")
    return 0 if ok else 1


if __name__ == "__main__":
    raise SystemExit(main())
