#!/usr/bin/env python3
"""Canonical technical-colossus proposal build/gate.

Purpose:
- Treat technical-colossus.md as the section contract.
- Treat MASTER_TEMPLATE_LOCKED.docx as the Word formatting contract.
- Keep technical-colossus.docx as the locked skeleton reference artifact.
- Fail fast when content does not follow the largest technical skeleton.

This script does not draft proposal prose. It gates and assembles already-drafted
Word-compatible Markdown through the one approved path.
"""
from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import subprocess
import sys
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
WRAPPER = ROOT / "bin" / "bid-manager"
BID = ROOT / "legacy" / "bid-manager"
SKELETON_MD = BID / "skeletons" / "1_technical" / "markdown" / "technical-colossus.md"
SKELETON_DOCX = BID / "skeletons" / "1_technical" / "docx" / "technical-colossus.docx"
MASTER_TEMPLATE = BID / "templates" / "MASTER_TEMPLATE_LOCKED.docx"
OUTPUT_ROOT = ROOT / "settlemint-office-agents" / "bid-manager" / "output"

INTERNAL_HEADING_PREFIXES = (
    "Technical Proposal Blueprint:",
    "Blueprint Rules:",
    "Structural Compliance",
    "Consistency Anchors",
    "Pre-Write Checklist",
    "Output Quality Gates",
    "MANDATORY DIAGRAMS",
    "Original Skeleton-Specific Rules",
    "Global Writer Guidance",
    "Proposal Objective",
    "Global Source Priority",
    "Global Visual System",
    "Visual Element Policy",
    "Section-to-Screenshot Map",
    "Appendix Guidance",
)
OPTIONAL_OR_INSTRUCTIONAL_CONTAINS = (
    " > Avoid",
    "Avoid",
    "Tone and Avoidance",
    "Visuals",
    "Writer's Checklist",
    "COLOSSUS",
    "Mermaid Coverage Seed",
    "Drafting Rule",
)
BANNED_VISIBLE_TERMS = [
    "PLACEHOLDER",
    "[VARIABLE",
    "TODO",
    "Logo Here",
    "COLOSSUS",
    "bid-manager",
    "skeleton",
    "proposal should",
    "client-ready",
    "buyer-safe",
    "internal review",
    "honest answer",
    "Writer's Checklist",
]


def slug(s: str) -> str:
    s = re.sub(r"[^a-z0-9]+", "-", s.lower()).strip("-")
    return s


def extract_headings(md: Path) -> list[dict]:
    out: list[dict] = []
    for line in md.read_text(encoding="utf-8", errors="ignore").splitlines():
        m = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if not m:
            continue
        text = re.sub(r"\s+", " ", m.group(2).strip())
        out.append({"level": len(m.group(1)), "text": text, "slug": slug(text)})
    return out


def required_contract() -> list[dict]:
    raw = extract_headings(SKELETON_MD)
    required: list[dict] = []
    proposal_started = False
    for h in raw:
        text = h["text"]
        if text == "Cover Page and Front Matter":
            proposal_started = True
        if not proposal_started:
            continue
        if any(text.startswith(prefix) for prefix in INTERNAL_HEADING_PREFIXES):
            continue
        if any(marker in text for marker in OPTIONAL_OR_INSTRUCTIONAL_CONTAINS):
            continue
        required.append(h)
    return required


def content_headings(md: Path) -> set[str]:
    return {h["slug"] for h in extract_headings(md)}


def render_pdf(docx: Path, out_dir: Path) -> Path | None:
    if not shutil_which("libreoffice"):
        return None
    out_dir.mkdir(parents=True, exist_ok=True)
    proc = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx)],
        text=True,
        capture_output=True,
        timeout=240,
    )
    if proc.returncode != 0:
        raise RuntimeError(proc.stdout + proc.stderr)
    return out_dir / (docx.stem + ".pdf")


def shutil_which(name: str) -> str | None:
    from shutil import which
    return which(name)


def pdf_pages(pdf: Path | None) -> int | None:
    if not pdf or not pdf.exists() or not shutil_which("pdfinfo"):
        return None
    out = subprocess.check_output(["pdfinfo", str(pdf)], text=True)
    m = re.search(r"Pages:\s+(\d+)", out)
    return int(m.group(1)) if m else None


def zip_bad(docx: Path) -> str | None:
    with zipfile.ZipFile(docx) as z:
        return z.testzip()


def scan_docx_text(docx: Path) -> str:
    from docx import Document
    d = Document(str(docx))
    bits = [p.text for p in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            bits.append(" | ".join(c.text for c in row.cells))
    return "\n".join(bits)


def build(args: argparse.Namespace) -> dict:
    content = Path(args.content_md).resolve()
    if not content.exists():
        raise FileNotFoundError(content)
    out_dir = Path(args.out_dir).resolve() if args.out_dir else OUTPUT_ROOT / f"technical-colossus-{slug(content.stem)}"
    out_dir.mkdir(parents=True, exist_ok=True)
    docx = Path(args.output_docx).resolve() if args.output_docx else out_dir / f"{slug(content.stem)}-technical-colossus.docx"
    report_path = out_dir / "technical-colossus-build-report.json"

    required = required_contract()
    present = content_headings(content)
    missing = [h for h in required if h["slug"] not in present]
    payload: dict = {
        "ok": False,
        "mode": "technical-colossus-build",
        "content_md": str(content),
        "output_docx": str(docx),
        "skeleton_markdown_contract": str(SKELETON_MD),
        "skeleton_docx_reference": str(SKELETON_DOCX),
        "master_template_format_contract": str(MASTER_TEMPLATE),
        "required_heading_count": len(required),
        "missing_heading_count": len(missing),
        "missing_headings": missing[:200],
    }
    if missing and not args.allow_missing:
        payload["failure"] = "content markdown does not satisfy technical-colossus heading contract"
        report_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
        return payload

    # Use the canonical wrapper so template validation, output root enforcement,
    # and recovered converter behavior stay in one place.
    proc = subprocess.run([str(WRAPPER), "md-to-docx", str(content), str(docx)], text=True, capture_output=True, timeout=args.timeout)
    payload["converter_exit_code"] = proc.returncode
    payload["converter_stdout_tail"] = proc.stdout[-4000:]
    payload["converter_stderr_tail"] = proc.stderr[-4000:]
    if proc.returncode != 0 or not docx.exists():
        payload["failure"] = "canonical md-to-docx conversion failed"
        report_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
        return payload

    # Work around wrapper raw-JSON normalization limits by applying cover fields
    # after canonical conversion, then saving the same DOCX.
    sys.path.insert(0, str(BID / "scripts"))
    from markdown_to_docx import _update_cover_page  # type: ignore
    from docx import Document
    cover_fields = {
        "company": args.company,
        "title": args.title,
        "subtitle": args.subtitle,
        "valid_until": args.valid_until,
        "contact_name": args.contact_name,
        "contact_title": args.contact_title,
        "contact_email": args.contact_email,
        "contact_phone": args.contact_phone,
    }
    cover_doc = Document(str(docx))
    _update_cover_page(cover_doc, cover_fields)
    cover_doc.save(str(docx))

    full_text = scan_docx_text(docx)
    banned_hits = {term: full_text.lower().count(term.lower()) for term in BANNED_VISIBLE_TERMS}
    pdf = render_pdf(docx, out_dir / "pdf") if args.render_pdf else None
    payload.update({
        "zip_bad": zip_bad(docx),
        "sha256": hashlib.sha256(docx.read_bytes()).hexdigest(),
        "word_estimate": len(re.findall(r"\b\w+\b", full_text)),
        "banned_hits": banned_hits,
        "pdf": str(pdf) if pdf else None,
        "pdf_pages": pdf_pages(pdf),
    })
    payload["ok"] = payload["zip_bad"] is None and not any(banned_hits.values()) and payload["missing_heading_count"] == 0
    if not payload["ok"]:
        payload["failure"] = "post-build gate failed"
    report_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    return payload


def contract_payload() -> dict:
    required = required_contract()
    return {
        "ok": True,
        "mode": "technical-colossus-contract",
        "skeleton_markdown_contract": str(SKELETON_MD),
        "skeleton_docx_reference": str(SKELETON_DOCX),
        "master_template_format_contract": str(MASTER_TEMPLATE),
        "required_heading_count": len(required),
        "required_headings": required,
        "rule": "Largest technical proposal builds must satisfy this heading contract before DOCX assembly. Do not fall back to generic proposal structures.",
    }


def main(argv=None) -> int:
    parser = argparse.ArgumentParser(prog="technical-colossus-build")
    parser.add_argument("--contract", action="store_true", help="Print the canonical section contract and exit")
    parser.add_argument("--content-md", help="Drafted proposal markdown to gate and assemble")
    parser.add_argument("--out-dir", help="Output directory under bid-manager output root")
    parser.add_argument("--output-docx", help="Output DOCX path")
    parser.add_argument("--allow-missing", action="store_true", help="Do not fail before conversion when headings are missing; still reports them")
    parser.add_argument("--no-render-pdf", dest="render_pdf", action="store_false", help="Skip LibreOffice PDF render")
    parser.add_argument("--timeout", type=int, default=300, help="Converter timeout seconds")
    parser.add_argument("--company", default="Emirates NBD Group", help="Cover company/client name")
    parser.add_argument("--title", default="Tokenization Platform Technical Proposal", help="Cover title")
    parser.add_argument("--subtitle", default="Technical response prepared by SettleMint", help="Cover subtitle")
    parser.add_argument("--valid-until", default="30 June 2026", help="Cover validity date")
    parser.add_argument("--contact-name", default="SettleMint", help="Cover contact name")
    parser.add_argument("--contact-title", default="Proposal Team", help="Cover contact title")
    parser.add_argument("--contact-email", default="info@settlemint.com", help="Cover contact email")
    parser.add_argument("--contact-phone", default="", help="Cover contact phone")
    args = parser.parse_args(argv)
    if args.contract or not args.content_md:
        print(json.dumps(contract_payload(), indent=2))
        return 0
    payload = build(args)
    print(json.dumps(payload, indent=2))
    return 0 if payload.get("ok") else 1


if __name__ == "__main__":
    raise SystemExit(main())
