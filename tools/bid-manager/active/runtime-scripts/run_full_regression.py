#!/usr/bin/env python3
"""Run full technical/commercial proposal conversion and render QA regression."""
from __future__ import annotations

import argparse
import json
import re
import subprocess
import time
import zipfile
from pathlib import Path
from typing import Any

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
WRAPPER = ROOT / "bin" / "bid-manager"
BID = ROOT / "legacy" / "bid-manager"

FULL_CASES = [
    {
        "name": "technical_full",
        "variant": "full",
        "source": BID / "skeletons" / "1_technical" / "markdown" / "technical-full.md",
        "docx_name": "bid-manager-technical-full-regression.docx",
    },
    {
        "name": "commercial_full",
        "variant": "full",
        "source": BID / "skeletons" / "2_commercial" / "markdown" / "commercial-full.md",
        "docx_name": "bid-manager-commercial-full-regression.docx",
    },
]

BANNED_ROUTING_NAMES = ["proposal-maker", "proposal-factory"]


def run_cmd(name: str, cmd: list[str], cwd: Path | None = None) -> dict[str, Any]:
    start = time.time()
    proc = subprocess.run(cmd, cwd=str(cwd or ROOT), text=True, capture_output=True)
    return {
        "name": name,
        "cmd": cmd,
        "exit": proc.returncode,
        "ok": proc.returncode == 0,
        "duration_seconds": round(time.time() - start, 3),
        "stdout_tail": proc.stdout[-4000:],
        "stderr_tail": proc.stderr[-4000:],
    }


def docx_gate(path: Path) -> dict[str, Any]:
    result: dict[str, Any] = {"path": str(path), "exists": path.exists(), "ok": False}
    if not path.exists():
        return result
    result["bytes"] = path.stat().st_size
    try:
        with zipfile.ZipFile(path) as zf:
            bad = zf.testzip()
            names = set(zf.namelist())
        result["zip_ok"] = bad is None
        result["has_document_xml"] = "word/document.xml" in names
    except Exception as exc:  # noqa: BLE001
        result["zip_ok"] = False
        result["zip_error"] = repr(exc)
        return result
    try:
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                text += "\n" + " | ".join(cell.text for cell in row.cells)
        result["paragraphs"] = len(doc.paragraphs)
        result["tables"] = len(doc.tables)
        result["word_count"] = len(re.findall(r"\b\w+\b", text))
        hits = [name for name in BANNED_ROUTING_NAMES if name in text.lower()]
        result["banned_routing_name_hits"] = hits
        result["docx_parse_ok"] = True
    except Exception as exc:  # noqa: BLE001
        result["docx_parse_ok"] = False
        result["docx_error"] = repr(exc)
        return result
    result["ok"] = bool(
        result.get("zip_ok")
        and result.get("has_document_xml")
        and result.get("docx_parse_ok")
        and result.get("word_count", 0) >= 1500
        and not result.get("banned_routing_name_hits")
    )
    return result


def render_pdf(docx_path: Path, out_dir: Path) -> tuple[dict[str, Any], Path]:
    cmd = ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)]
    case = run_cmd(f"render_pdf_{docx_path.stem}", cmd)
    pdf = out_dir / (docx_path.stem + ".pdf")
    case["pdf"] = str(pdf)
    case["pdf_exists"] = pdf.exists()
    return case, pdf


def pdf_gate(pdf_path: Path) -> dict[str, Any]:
    result: dict[str, Any] = {"path": str(pdf_path), "exists": pdf_path.exists(), "ok": False}
    if not pdf_path.exists():
        return result
    info = run_cmd(f"pdfinfo_{pdf_path.stem}", ["pdfinfo", str(pdf_path)])
    result["pdfinfo"] = info
    pages = None
    for line in info["stdout_tail"].splitlines():
        if line.startswith("Pages:"):
            try:
                pages = int(line.split(":", 1)[1].strip())
            except ValueError:
                pass
    result["pages"] = pages
    text_case = run_cmd(f"pdftotext_{pdf_path.stem}", ["pdftotext", str(pdf_path), "-"])
    result["pdftotext"] = {k: v for k, v in text_case.items() if k not in {"stdout_tail"}}
    text = text_case["stdout_tail"]
    result["text_tail_chars"] = len(text)
    result["contains_bid_manager"] = "bid-manager" in text.lower() or "bid manager" in text.lower()
    result["ok"] = bool(info["ok"] and text_case["ok"] and pages and pages >= 3 and len(text) > 500)
    return result


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Run bid-manager full proposal conversion/render regression")
    parser.add_argument("--json", action="store_true")
    parser.add_argument("--out-dir", default=None)
    ns = parser.parse_args(argv)

    stamp = time.strftime("%Y%m%dT%H%M%SZ", time.gmtime())
    out_dir = Path(ns.out_dir).resolve() if ns.out_dir else BID / "output" / f"full-regression-{stamp}"
    out_dir.mkdir(parents=True, exist_ok=True)

    cases = []
    gates = []
    for item in FULL_CASES:
        source = item["source"]
        docx_path = out_dir / item["docx_name"]
        convert = run_cmd(f"{item['name']}_md_to_docx", [str(WRAPPER), "md-to-docx", str(source), str(docx_path)])
        cases.append(convert)
        dg = docx_gate(docx_path)
        gates.append({"name": f"{item['name']}_docx_gate", **dg})
        if convert["ok"] and dg["ok"]:
            render, pdf_path = render_pdf(docx_path, out_dir)
            cases.append(render)
            gates.append({"name": f"{item['name']}_pdf_gate", **pdf_gate(pdf_path)})

    unexpected_cases = [c for c in cases if not c.get("ok")]
    failed_gates = [g for g in gates if not g.get("ok")]
    payload = {
        "agent": "bid-manager",
        "root": str(ROOT),
        "artifact_dir": str(out_dir),
        "case_count": len(cases),
        "passed_count": sum(1 for c in cases if c.get("ok")),
        "gate_count": len(gates),
        "gate_passed_count": sum(1 for g in gates if g.get("ok")),
        "unexpected_cases": unexpected_cases,
        "failed_gates": failed_gates,
        "cases": cases,
        "gates": gates,
        "ok": not unexpected_cases and not failed_gates,
    }
    (out_dir / "full-regression.json").write_text(json.dumps(payload, indent=2), encoding="utf-8")
    if ns.json:
        print(json.dumps(payload, indent=2))
    else:
        print(f"bid-manager full regression: {payload['passed_count']}/{payload['case_count']} commands passed; {payload['gate_passed_count']}/{payload['gate_count']} gates passed")
        print(f"artifacts: {out_dir}")
        for gate in gates:
            print(f"- {'OK' if gate.get('ok') else 'FAIL'}: {gate['name']}")
    return 0 if payload["ok"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
