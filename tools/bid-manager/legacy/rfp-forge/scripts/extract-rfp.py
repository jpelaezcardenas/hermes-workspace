#!/usr/bin/env python3
"""Extract questions/requirements from an RFP document.
Supports: XLSX, DOCX, PDF
Output: JSON with extracted questions and document structure metadata.

Usage: python3 extract-rfp.py <input_file>
"""
import json, sys, os, re

def extract_xlsx(path):
    """Extract from Excel - look for question columns and response columns."""
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    questions = []
    metadata = {"format": "xlsx", "sheets": []}
    
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=False))
        if not rows:
            continue
        
        # Find header row - look for keywords
        header_row_idx = None
        question_col = None
        response_col = None
        q_keywords = ['question', 'requirement', 'criteria', 'description', 'item', 'query', 'ask']
        r_keywords = ['response', 'answer', 'vendor', 'reply', 'comment', 'settlemint']
        
        for i, row in enumerate(rows[:10]):  # Check first 10 rows for headers
            cells = [str(c.value or '').lower().strip() for c in row]
            for j, cell in enumerate(cells):
                if any(k in cell for k in q_keywords) and question_col is None:
                    question_col = j
                    header_row_idx = i
                if any(k in cell for k in r_keywords) and response_col is None:
                    response_col = j
        
        # If no clear header, assume column A = questions, find first empty column for responses
        if question_col is None:
            # Look for the column with the most text content
            col_lengths = {}
            for row in rows[1:min(20, len(rows))]:
                for j, cell in enumerate(row):
                    v = str(cell.value or '').strip()
                    if v:
                        col_lengths[j] = col_lengths.get(j, 0) + len(v)
            if col_lengths:
                question_col = max(col_lengths, key=col_lengths.get)
                header_row_idx = 0
        
        if question_col is None:
            continue
        
        sheet_meta = {
            "name": sheet_name,
            "header_row": header_row_idx,
            "question_col": question_col,
            "response_col": response_col,
            "total_rows": len(rows)
        }
        metadata["sheets"].append(sheet_meta)
        
        # Extract questions
        start = (header_row_idx + 1) if header_row_idx is not None else 1
        for i, row in enumerate(rows[start:], start=start):
            cells = [c.value for c in row]
            if question_col < len(cells):
                q_text = str(cells[question_col] or '').strip()
                if len(q_text) > 5:  # Skip very short/empty cells
                    # Get row number (reference for writing back)
                    row_num = str(row[0].row) if hasattr(row[0], 'row') else str(i + 1)
                    
                    # Get existing response if any
                    existing = ''
                    if response_col is not None and response_col < len(cells):
                        existing = str(cells[response_col] or '').strip()
                    
                    # Get any ID/number from earlier columns
                    q_id = ''
                    if question_col > 0 and len(cells) > 0:
                        q_id = str(cells[0] or '').strip()
                    
                    questions.append({
                        "id": q_id,
                        "question": q_text,
                        "existing_response": existing,
                        "sheet": sheet_name,
                        "row": int(row_num),
                        "col": question_col
                    })
    
    wb.close()
    return {"questions": questions, "metadata": metadata}


def extract_docx(path):
    """Extract from Word - identify questions by numbering, bold text, or table structure."""
    import docx
    doc = docx.Document(path)
    questions = []
    metadata = {"format": "docx", "sections": []}
    
    # Check tables first (many RFPs use tables)
    for t_idx, table in enumerate(doc.tables):
        header_cells = [c.text.lower().strip() for c in table.rows[0].cells] if table.rows else []
        q_col = None
        r_col = None
        q_keywords = ['question', 'requirement', 'criteria', 'description', 'item']
        r_keywords = ['response', 'answer', 'vendor', 'reply']
        
        for j, cell in enumerate(header_cells):
            if any(k in cell for k in q_keywords):
                q_col = j
            if any(k in cell for k in r_keywords):
                r_col = j
        
        if q_col is not None:
            for r_idx, row in enumerate(table.rows[1:], start=1):
                cells = [c.text.strip() for c in row.cells]
                if q_col < len(cells) and len(cells[q_col]) > 5:
                    existing = cells[r_col] if r_col is not None and r_col < len(cells) else ''
                    q_id = cells[0] if q_col > 0 else ''
                    questions.append({
                        "id": q_id,
                        "question": cells[q_col],
                        "existing_response": existing,
                        "table": t_idx,
                        "row": r_idx,
                        "col": q_col
                    })
            metadata["sections"].append({"type": "table", "index": t_idx, "rows": len(table.rows)})
    
    # Also extract numbered paragraphs as potential questions
    current_section = ""
    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # Section headers
        if para.style.name.startswith('Heading'):
            current_section = text
            continue
        
        # Numbered items or questions (lines ending with ?)
        is_numbered = bool(re.match(r'^[\d]+[.\)]\s', text))
        is_question = text.endswith('?')
        is_requirement = any(k in text.lower() for k in ['shall', 'must', 'should', 'please describe', 'please provide', 'explain how'])
        
        if (is_numbered or is_question or is_requirement) and len(text) > 10:
            questions.append({
                "id": re.match(r'^([\d]+[.\)])', text).group(1) if is_numbered else "",
                "question": text,
                "existing_response": "",
                "section": current_section,
                "paragraph": p_idx
            })
    
    return {"questions": questions, "metadata": metadata}


def extract_pdf(path):
    """Extract from PDF - convert to text and identify questions."""
    import pdfplumber
    questions = []
    metadata = {"format": "pdf", "pages": 0}
    
    with pdfplumber.open(path) as pdf:
        metadata["pages"] = len(pdf.pages)
        full_text = ""
        
        for page in pdf.pages:
            # Try tables first
            tables = page.extract_tables()
            for t_idx, table in enumerate(tables):
                if not table or len(table) < 2:
                    continue
                header = [str(c or '').lower().strip() for c in table[0]]
                q_col = None
                r_col = None
                for j, h in enumerate(header):
                    if any(k in h for k in ['question', 'requirement', 'criteria', 'description']):
                        q_col = j
                    if any(k in h for k in ['response', 'answer', 'vendor']):
                        r_col = j
                
                if q_col is not None:
                    for row in table[1:]:
                        if q_col < len(row) and row[q_col] and len(str(row[q_col]).strip()) > 5:
                            existing = str(row[r_col] or '').strip() if r_col and r_col < len(row) else ''
                            questions.append({
                                "id": str(row[0] or '').strip() if q_col > 0 else '',
                                "question": str(row[q_col]).strip(),
                                "existing_response": existing,
                                "page": page.page_number
                            })
            
            # Also get free text
            text = page.extract_text() or ""
            full_text += text + "\n"
        
        # If no table questions found, parse text for questions
        if not questions:
            for line in full_text.split('\n'):
                line = line.strip()
                is_numbered = bool(re.match(r'^[\d]+[.\)]\s', line))
                is_question = line.endswith('?')
                is_requirement = any(k in line.lower() for k in ['shall', 'must', 'should', 'please describe', 'please provide'])
                
                if (is_numbered or is_question or is_requirement) and len(line) > 15:
                    questions.append({
                        "id": re.match(r'^([\d]+[.\)])', line).group(1) if is_numbered else "",
                        "question": line,
                        "existing_response": ""
                    })
    
    return {"questions": questions, "metadata": metadata}


def main():
    if len(sys.argv) < 2:
        print(json.dumps({"error": "Usage: extract-rfp.py <input_file>"}))
        sys.exit(1)
    
    path = sys.argv[1]
    if not os.path.exists(path):
        print(json.dumps({"error": f"File not found: {path}"}))
        sys.exit(1)
    
    ext = os.path.splitext(path)[1].lower()
    
    if ext in ['.xlsx', '.xls']:
        result = extract_xlsx(path)
    elif ext in ['.docx', '.doc']:
        result = extract_docx(path)
    elif ext == '.pdf':
        result = extract_pdf(path)
    else:
        print(json.dumps({"error": f"Unsupported format: {ext}. Supported: xlsx, docx, pdf"}))
        sys.exit(1)
    
    result["source_file"] = path
    result["source_format"] = ext
    result["total_questions"] = len(result["questions"])
    
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()
