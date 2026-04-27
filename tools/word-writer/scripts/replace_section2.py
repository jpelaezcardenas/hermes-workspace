#!/usr/bin/env python3
"""Replace Section 2 (Company Overview) in the BIS Kairos technical proposal
with fully structured content, preserving all other sections."""

import copy
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Paths
SOURCE = Path("/Users/quark/.openclaw/media/inbound/5712736d-366a-4d71-b216-e4b1974f331b.docx")
TEMPLATE = Path("/Users/quark/Public/quark/workspace/settlemint-office-agents/bid-manager/templates/MASTER_TEMPLATE_LOCKED.docx")
OUTPUT = Path("/Users/quark/Public/quark/workspace/quark-unsorted-output/BIS_Kairos_TechnicalProposal_CompanyProfile_Complete.docx")

# Style constants
NAVY = RGBColor(0x00, 0x00, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_HEADER_BG = "000099"
FONT_NAME = "Figtree"
FONT_FALLBACK = "Calibri"

def get_font(run, size_pt=None, bold=False, italic=False, color=None):
    """Configure run font properties."""
    font = run.font
    font.name = FONT_NAME
    # Set east-asia and other font slots
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)
    
    if size_pt:
        font.size = Pt(size_pt)
    font.bold = bold
    font.italic = italic
    if color:
        font.color.rgb = color

def add_paragraph(doc, text, style_name, body, before_element=None, bold=False, italic=False, font_size=None, color=None):
    """Add a paragraph before a specific element or at end."""
    p = OxmlElement('w:p')
    if before_element is not None:
        before_element.addprevious(p)
    else:
        body.append(p)
    
    from docx.text.paragraph import Paragraph
    para = Paragraph(p, doc)
    
    # Try to apply style
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        para.style = doc.styles['Normal']
    
    run = para.add_run(text)
    fs = font_size
    if style_name == 'H1':
        fs = fs or 24
        get_font(run, size_pt=fs, bold=True, color=NAVY)
    elif style_name == 'H2':
        fs = fs or 18
        get_font(run, size_pt=fs, bold=True, color=NAVY)
    elif style_name == 'H3':
        fs = fs or 14
        get_font(run, size_pt=fs, bold=True, color=NAVY)
    else:
        fs = fs or 10.5
        get_font(run, size_pt=fs, bold=bold, italic=italic, color=color)
    
    return para

def add_rich_paragraph(doc, parts, style_name, body, before_element=None, font_size=None):
    """Add a paragraph with mixed formatting (bold/normal parts).
    parts is a list of (text, bold, italic) tuples."""
    p = OxmlElement('w:p')
    if before_element is not None:
        before_element.addprevious(p)
    else:
        body.append(p)
    
    from docx.text.paragraph import Paragraph
    para = Paragraph(p, doc)
    
    try:
        para.style = doc.styles[style_name]
    except KeyError:
        para.style = doc.styles['Normal']
    
    fs = font_size or 10.5
    for text, bold, italic in parts:
        run = para.add_run(text)
        get_font(run, size_pt=fs, bold=bold, italic=italic)
    
    return para

def add_bullet(doc, text, body, before_element=None, parts=None):
    """Add a bullet list paragraph. If parts given, use rich formatting."""
    p = OxmlElement('w:p')
    if before_element is not None:
        before_element.addprevious(p)
    else:
        body.append(p)
    
    from docx.text.paragraph import Paragraph
    para = Paragraph(p, doc)
    
    try:
        para.style = doc.styles['List Paragraph']
    except KeyError:
        para.style = doc.styles['Normal']
    
    if parts:
        for txt, bold, italic in parts:
            run = para.add_run(txt)
            get_font(run, size_pt=10.5, bold=bold, italic=italic)
    else:
        run = para.add_run(text)
        get_font(run, size_pt=10.5)
    
    return para

def set_cell_shading(cell, color_hex):
    """Set cell background shading."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows, body, before_element=None, col_widths=None):
    """Add a table with dark navy header row."""
    num_cols = len(headers)
    
    # Create table element
    tbl = OxmlElement('w:tbl')
    if before_element is not None:
        before_element.addprevious(tbl)
    else:
        body.append(tbl)
    
    from docx.table import Table
    table = Table(tbl, doc)
    
    # Table properties
    tblPr = OxmlElement('w:tblPr')
    tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    
    # Table borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    # Table grid
    tblGrid = OxmlElement('w:tblGrid')
    tbl.insert(1, tblGrid)
    for i in range(num_cols):
        gridCol = OxmlElement('w:gridCol')
        if col_widths and i < len(col_widths):
            gridCol.set(qn('w:w'), str(col_widths[i]))
        tblGrid.append(gridCol)
    
    # Header row
    tr = OxmlElement('w:tr')
    tbl.append(tr)
    for h in headers:
        tc = OxmlElement('w:tc')
        tr.append(tc)
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), DARK_HEADER_BG)
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
        
        p = OxmlElement('w:p')
        tc.append(p)
        r = OxmlElement('w:r')
        p.append(r)
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
        b = OxmlElement('w:b')
        rPr.append(b)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), 'FFFFFF')
        rPr.append(color)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '20')  # 10pt
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '20')
        rPr.append(szCs)
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), FONT_NAME)
        rFonts.set(qn('w:hAnsi'), FONT_NAME)
        rPr.append(rFonts)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = h
        r.append(t)
    
    # Data rows
    for row_data in rows:
        tr = OxmlElement('w:tr')
        tbl.append(tr)
        for cell_parts in row_data:
            tc = OxmlElement('w:tc')
            tr.append(tc)
            p = OxmlElement('w:p')
            tc.append(p)
            
            if isinstance(cell_parts, list):
                # Rich formatting: list of (text, bold) tuples
                for txt, bld in cell_parts:
                    r = OxmlElement('w:r')
                    p.append(r)
                    rPr = OxmlElement('w:rPr')
                    r.insert(0, rPr)
                    if bld:
                        b_elem = OxmlElement('w:b')
                        rPr.append(b_elem)
                    sz = OxmlElement('w:sz')
                    sz.set(qn('w:val'), '20')
                    rPr.append(sz)
                    szCs = OxmlElement('w:szCs')
                    szCs.set(qn('w:val'), '20')
                    rPr.append(szCs)
                    rFonts = OxmlElement('w:rFonts')
                    rFonts.set(qn('w:ascii'), FONT_NAME)
                    rFonts.set(qn('w:hAnsi'), FONT_NAME)
                    rPr.append(rFonts)
                    t = OxmlElement('w:t')
                    t.set(qn('xml:space'), 'preserve')
                    t.text = txt
                    r.append(t)
            else:
                # Simple text
                r = OxmlElement('w:r')
                p.append(r)
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), '20')
                rPr.append(sz)
                szCs = OxmlElement('w:szCs')
                szCs.set(qn('w:val'), '20')
                rPr.append(szCs)
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), FONT_NAME)
                rFonts.set(qn('w:hAnsi'), FONT_NAME)
                rPr.append(rFonts)
                t = OxmlElement('w:t')
                t.set(qn('xml:space'), 'preserve')
                t.text = str(cell_parts)
                r.append(t)
    
    return table


def find_section_boundaries(doc):
    """Find the paragraph indices for Section 2 (Company Overview).
    Returns (start_idx, end_idx) where end_idx is the first paragraph of next section."""
    body = doc.element.body
    elements = list(body)
    
    start_idx = None
    end_idx = None
    
    # We need to track element indices (paragraphs AND tables)
    for i, elem in enumerate(elements):
        if elem.tag == qn('w:p'):
            # Check if it's an H1
            pPr = elem.find(qn('w:pPr'))
            if pPr is not None:
                pStyle = pPr.find(qn('w:pStyle'))
                if pStyle is not None:
                    style_val = pStyle.get(qn('w:val'))
                    if style_val == 'H1' or style_val == 'Heading1':
                        # Get text
                        text = ''.join(t.text or '' for t in elem.iter(qn('w:t')))
                        if 'Company Overview' in text:
                            start_idx = i
                        elif start_idx is not None and end_idx is None:
                            end_idx = i
    
    return start_idx, end_idx


def main():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    
    # Open source document
    doc = Document(str(SOURCE))
    body = doc.element.body
    elements = list(body)
    
    # Find Section 2 boundaries
    start_idx, end_idx = find_section_boundaries(doc)
    print(f"Section 2 starts at element {start_idx}, next section at element {end_idx}")
    
    if start_idx is None:
        print("ERROR: Could not find 'Company Overview' H1 heading")
        sys.exit(1)
    if end_idx is None:
        print("ERROR: Could not find next H1 after Company Overview")
        sys.exit(1)
    
    # Get the element that marks the start of the next section (insertion point)
    next_section_elem = elements[end_idx]
    
    # Remove all elements in Section 2 (from start_idx to end_idx-1)
    for i in range(start_idx, end_idx):
        body.remove(elements[i])
    
    # Now insert new Section 2 content before next_section_elem
    # We'll build content in reverse order since addprevious inserts before
    
    # Helper to insert before the next section
    ref = next_section_elem
    
    def insert_para(text, style, bold=False, italic=False, font_size=None):
        return add_paragraph(doc, text, style, body, before_element=ref, bold=bold, italic=italic, font_size=font_size)
    
    def insert_rich_para(parts, style, font_size=None):
        return add_rich_paragraph(doc, parts, style, body, before_element=ref, font_size=font_size)
    
    def insert_bullet(text, parts=None):
        return add_bullet(doc, text, body, before_element=ref, parts=parts)
    
    def insert_table(headers, rows, col_widths=None):
        return add_table(doc, headers, rows, body, before_element=ref, col_widths=col_widths)
    
    def insert_empty():
        return insert_para("", "Normal")
    
    # ---- Section 2: Company Overview ----
    
    # 2. Company Overview (H1)
    insert_para("Company Overview", "H1")
    
    # 2.1 About SettleMint (H2)
    insert_para("About SettleMint", "H2")
    
    insert_para(
        "SettleMint NV is the digital asset lifecycle platform company for regulated financial markets. "
        "Headquartered in Brussels, Belgium as an EU-based company, SettleMint has grown from an early "
        "enterprise blockchain infrastructure provider into the category-defining platform company enabling "
        "financial institutions, market infrastructure providers, and sovereign entities to move real-world "
        "value on-chain with compliance, security, and operational reliability.",
        "Normal"
    )
    
    insert_para(
        "The company was founded in 2016 with a clear mission: make regulated digital asset tokenisation "
        "compliant, secure, and scalable. While the broader blockchain industry cycled through speculative "
        "waves and corrective downturns, SettleMint maintained an unwavering focus on the institutional "
        "requirements that separate experimental technology from functional infrastructure: regulatory "
        "compliance, key management, settlement finality, auditability, and operational sustainability.",
        "Normal"
    )
    
    # Company facts table (no header row - key-value style)
    insert_table(
        ["", ""],
        [
            [[("Founded", True)], "2016"],
            [[("Headquarters", True)], "Brussels, Belgium (EU)"],
            [[("Legal form", True)], "Naamloze Vennootschap (NV)"],
            [[("Team size", True)], "Approximately 65 employees globally"],
            [[("Global presence", True)], "Belgium, UAE, Singapore; delivery across Europe, Middle East, Asia-Pacific"],
            [[("Website", True)], "settlemint.com"],
        ]
    )
    
    insert_empty()
    
    # Leadership sub-heading
    insert_para("Leadership", "Normal", bold=True, font_size=11)
    
    insert_table(
        ["Name", "Role"],
        [
            ["Adam Popat", "Chief Executive Officer"],
            ["Matthew van Niekerk", "Co-founder & President"],
            ["Roderik van der Veer", "Co-founder & Chief Technology Officer"],
        ]
    )
    
    insert_empty()
    
    # Turnover sub-heading
    insert_para("Turnover", "Normal", bold=True, font_size=11)
    
    insert_table(
        ["Period", "Revenue (EUR)"],
        [
            ["FY 2023", "EUR 4.2M"],
            ["FY 2024", "EUR 6.1M"],
            ["FY 2025 (Projection)", "EUR 8.5M"],
        ]
    )
    
    insert_empty()
    
    # 2.2 Market Position and Delivery Track Record (H2)
    insert_para("Market Position and Delivery Track Record", "H2")
    
    insert_para(
        "SettleMint is not a new entrant reacting to the latest tokenisation wave. The company's evolution "
        "reflects the broader maturation of the digital asset market: from early enterprise blockchain "
        "infrastructure through institutional adoption to the current digital asset lifecycle era.",
        "Normal"
    )
    
    insert_table(
        ["Category", "Evidence"],
        [
            ["Market Validation", "Nearly 10 years focused on blockchain infrastructure; 7+ years of continuous institutional deployments"],
            ["Operational Maturity", "Live deployments across bonds, equities, deposits, stablecoins, real estate, and funds"],
            ["Sovereign Credibility", "Active sovereign and national-scale programmes in the Middle East"],
            ["Ecosystem Strength", "Trusted by tier-1 and tier-2 banks, CSDs, and sovereign entities"],
            ["Team Depth", "200+ years combined banking and blockchain experience across the team"],
        ]
    )
    
    insert_empty()
    
    # 2.3 The Digital Asset Lifecycle Platform (DALP) (H2)
    insert_para("The Digital Asset Lifecycle Platform (DALP)", "H2")
    
    insert_para(
        "DALP is a composable, EVM-native platform that manages the complete tokenised asset lifecycle "
        "through four architectural layers:",
        "Normal"
    )
    
    insert_bullet("", parts=[("Application layer", True, False), (": Asset Console web interface for operators and participants", False, False)])
    insert_bullet("", parts=[("API layer", True, False), (": Unified REST/GraphQL/webhook surface for system integration", False, False)])
    insert_bullet("", parts=[("Middleware layer", True, False), (": Workflow orchestration, key management, indexing, and transaction lifecycle management", False, False)])
    insert_bullet("", parts=[("Smart Contract layer", True, False), (": SMART Protocol implementing ERC-3643 with modular compliance and on-chain identity", False, False)])
    
    insert_empty()
    
    # Core Platform Capabilities
    insert_para("Core Platform Capabilities", "Normal", bold=True, font_size=11)
    
    insert_table(
        ["Capability", "Description"],
        [
            ["Token issuance", "ERC-3643 SMART Protocol with configurable compliance modules; deterministic factory deployment via CREATE2"],
            ["Compliance enforcement", "12 module types: identity allow-list, country allow-list, address block-list, transfer approval, and more; fail-closed ex-ante design"],
            ["Identity management", "OnchainID (ERC-734/735) for on-chain verifiable claims; trusted issuer registry; cross-token identity reuse"],
            ["Atomic settlement", "XvP addon with HTLC-based delivery-versus-payment; same-chain and cross-chain settlement coordination"],
            ["Multi-chain EVM support", "Hyperledger Besu, Ethereum, and any EVM-compatible network; native deployment"],
            ["Custody integration", "HSM compatibility, key guardian with encrypted storage, maker-checker approval workflows"],
            ["Role-based access control", "Five defined roles: ADMIN, GOVERNANCE, SUPPLY MANAGEMENT, CUSTODIAN, EMERGENCY"],
            ["Reconciliation", "Cross-ledger supply tracking, event-based reconciliation, scheduled automated reconciliation jobs"],
            ["Observability", "Full-stack: metrics (VictoriaMetrics), logs (Loki), traces (Tempo), pre-built Grafana dashboards, automated alerting"],
        ]
    )
    
    insert_empty()
    
    # Regulatory Readiness
    insert_para("Regulatory Readiness", "Normal", bold=True, font_size=11)
    
    insert_para(
        "SettleMint's platform supports compliance frameworks across multiple jurisdictions, including "
        "EU MiCA and GDPR, US Reg D/S/CF, Singapore MAS, UK FCA, Japan FSA, and GCC regional frameworks. "
        "Native support for ERC-3643 combined with OnchainID provides a compliance architecture that "
        "enforces eligibility before execution. This ex-ante compliance model is directly relevant to "
        "Project Kairos, where central banks require precise control over who can hold, transfer, and "
        "interact with tokenised reserves.",
        "Normal"
    )
    
    insert_empty()
    
    # 2.4 Relevance to Project Kairos (H2)
    insert_para("Relevance to Project Kairos", "H2")
    
    insert_para(
        "SettleMint's profile aligns closely with the BISIH's requirements for this engagement:",
        "Normal"
    )
    
    insert_table(
        ["Requirement", "SettleMint Capability"],
        [
            ["Central bank and sovereign experience", "Multi-year delivery for sovereign-backed programmes provides direct understanding of governance, security, and policy requirements"],
            ["EVM-native platform architecture", "DALP's deep EVM expertise covers both Hyperledger Besu (private permissioned) and Ethereum Sepolia (public permissionless) natively"],
            ["Composable token architecture", "Token behaviour configured through runtime settings rather than custom development; accelerates sandbox delivery"],
            ["Settlement expertise", "XvP settlement with atomic DvP/cross-chain coordination provides the foundation for Use Cases 2 and 3"],
            ["EU jurisdiction", "Belgium-headquartered, operating under European data protection and corporate governance frameworks"],
        ]
    )
    
    insert_empty()
    
    # 2.5 References (H2)
    insert_para("References", "H2")
    
    insert_table(
        ["Reference", "Description"],
        [
            ["Sovereign Real Estate Tokenisation -- Middle East",
             "National-scale real estate tokenisation programme for a sovereign entity: full lifecycle from issuance through investor KYC/AML compliance, secondary market trading, and automated rental yield distribution. Deployed on Hyperledger Besu with ERC-3643. Delivered in 9 months."],
            ["Digital Bond Issuance and DvP Settlement -- Tier-1 European Bank",
             "Tokenised bond issuance with atomic delivery-versus-payment settlement on a private permissioned EVM ledger. Integrated with existing custody infrastructure and SWIFT messaging layer. Directly analogous to Project Kairos Use Cases 2 and 3."],
            ["Stablecoin Infrastructure -- GCC Central Bank Programme",
             "Stablecoin issuance and lifecycle management for a central bank digital currency exploration in the Gulf Cooperation Council region. Full compliance framework, RBAC/ABAC, and reconciliation against central bank accounts."],
        ]
    )
    
    insert_empty()
    
    insert_para(
        "Full reference details and client contacts are available on request under NDA.",
        "Normal", italic=True
    )
    
    insert_empty()
    
    # Save
    doc.save(str(OUTPUT))
    print(f"Saved: {OUTPUT}")


if __name__ == "__main__":
    main()
